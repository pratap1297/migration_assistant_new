"""
Word Document Generator for HLD Documents
Converts markdown HLD content to professional Word documents
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from src.core.models import RepositoryAnalysis, HLDContent


class WordDocumentGenerator:
    """Generates professional Word documents from HLD content"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word document generation. Install with: pip install python-docx")
        
        self.doc = None
        self.styles = {}
        
    def generate_hld_word_document(self, analysis: RepositoryAnalysis, 
                                 hld_content: HLDContent, 
                                 markdown_content: str,
                                 output_path: str) -> str:
        """Generate Word document from HLD content"""
        
        # Create new document
        self.doc = Document()
        
        # Setup document styles
        self._setup_document_styles()
        
        # Add document properties
        self._add_document_properties(analysis)
        
        # Parse markdown and convert to Word
        self._convert_markdown_to_word(markdown_content)
        
        # Save document
        self.doc.save(output_path)
        
        return output_path
    
    def _setup_document_styles(self):
        """Setup professional document styles"""
        
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading 1 style
        h1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(18)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
        
        # Heading 2 style
        h2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(47, 84, 150)  # Medium blue
        h2_style.paragraph_format.space_before = Pt(10)
        h2_style.paragraph_format.space_after = Pt(4)
        
        # Heading 3 style
        h3_style = self.doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(68, 114, 196)  # Light blue
        h3_style.paragraph_format.space_before = Pt(8)
        h3_style.paragraph_format.space_after = Pt(3)
        
        # Normal text style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        
        # Code style
        code_style = self.doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Store styles for reference
        self.styles = {
            'title': title_style,
            'h1': h1_style,
            'h2': h2_style,
            'h3': h3_style,
            'normal': normal_style,
            'code': code_style
        }
    
    def _add_document_properties(self, analysis: RepositoryAnalysis):
        """Add document properties and metadata"""
        
        # Set document properties
        core_props = self.doc.core_properties
        repo_name = getattr(analysis, 'repo_name', None)
        if not repo_name and hasattr(analysis, 'repository_url'):
            repo_name = analysis.repository_url.rstrip('/').split('/')[-1]
        if not repo_name:
            repo_name = 'Unknown Repository'
            
        core_props.title = f"High-Level Design - {repo_name} Azure Migration"
        core_props.author = "Migration Team"
        core_props.subject = "Azure Migration High-Level Design"
        core_props.comments = f"HLD document for {repo_name} migration to Azure"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
    
    def _convert_markdown_to_word(self, markdown_content: str):
        """Convert markdown content to Word document"""
        
        lines = markdown_content.split('\n')
        in_code_block = False
        in_table = False
        table_data = []
        table_headers = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    in_code_block = False
                    continue
                else:
                    # Start code block
                    in_code_block = True
                    continue
            
            if in_code_block:
                # Add code line
                p = self.doc.add_paragraph(line, style=self.styles['code'])
                continue
            
            # Handle tables
            if line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                    table_headers = []
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # Check if this is a separator line (contains only dashes)
                if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                    continue
                
                # Check if this is the header row
                if not table_headers:
                    table_headers = cells
                else:
                    table_data.append(cells)
                
                continue
            else:
                # End table if we were in one
                if in_table and table_data:
                    self._add_table(table_headers, table_data)
                    in_table = False
                    table_data = []
                    table_headers = []
            
            # Handle headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1).style = self.styles['h1']
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2).style = self.styles['h2']
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3).style = self.styles['h3']
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], level=4).style = self.styles['h3']
            elif line.startswith('- '):
                # Bullet list
                p = self.doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                # Numbered list
                p = self.doc.add_paragraph(line[3:], style='List Number')
            elif line.strip():
                # Regular paragraph with formatting
                self._add_formatted_paragraph(line)
            else:
                # Empty line - add spacing
                self.doc.add_paragraph()
        
        # Handle any remaining table
        if in_table and table_data:
            self._add_table(table_headers, table_data)
    
    def _add_formatted_paragraph(self, text: str):
        """Add paragraph with inline formatting"""
        p = self.doc.add_paragraph()
        
        # Split by formatting markers
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                # Italic text
                run = p.add_run(part[1:-1])
                run.font.italic = True
            else:
                # Regular text
                p.add_run(part)
    
    def _add_table(self, headers: List[str], data: List[List[str]]):
        """Add formatted table to document"""
        if not data:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            # Add light blue background
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>'))
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data):
                cell = table_row.cells[col_idx]
                cell.text = cell_data
        
        # Add spacing after table
        self.doc.add_paragraph()
    
    def _add_page_break(self):
        """Add page break"""
        self.doc.add_page_break()
    
    def _add_horizontal_rule(self):
        """Add horizontal rule"""
        p = self.doc.add_paragraph()
        p.add_run('_' * 80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_hld_word_document(analysis: RepositoryAnalysis, 
                             hld_content: HLDContent,
                             markdown_content: str,
                             output_dir: str) -> str:
    """Convenience function to generate HLD Word document"""
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename
    repo_name = getattr(analysis, 'repo_name', None)
    if not repo_name and hasattr(analysis, 'repository_url'):
        repo_name = analysis.repository_url.rstrip('/').split('/')[-1]
    if not repo_name:
        repo_name = 'Unknown_Repository'
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"HLD_{repo_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    
    # Generate document
    generator = WordDocumentGenerator()
    result_path = generator.generate_hld_word_document(
        analysis, hld_content, markdown_content, output_path
    )
    
    return result_path


if __name__ == "__main__":
    # Test the Word document generator
    print("Word Document Generator for HLD Documents")
    print("This module provides Word document generation capabilities for HLD documents.") 