"""
Professional Word Document Generator for Application Intelligence Reports
Supports custom templates and professional formatting
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    from matplotlib.patches import FancyBboxPatch
    import numpy as np
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    MATPLOTLIB_AVAILABLE = False

class ProfessionalWordGenerator:
    """
    Professional Word document generator with template support
    """
    
    def __init__(self, template_path: Optional[str] = None):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required. Install with: pip install python-docx")
        
        # Import Document here to avoid import errors
        from docx import Document
        
        self.template_path = template_path
        self.colors = {
            'primary': RGBColor(0, 120, 215),      # Professional blue
            'secondary': RGBColor(72, 72, 72),     # Dark gray
            'accent': RGBColor(0, 178, 148),       # Teal
            'warning': RGBColor(255, 140, 0),      # Orange
            'danger': RGBColor(232, 17, 35),       # Red
            'success': RGBColor(16, 124, 16),      # Green
            'light_gray': RGBColor(248, 249, 250), # Light background
            'border': RGBColor(222, 226, 230)      # Border color
        }
        
        # Default template structure
        self.template_structure = {
            'cover_page': True,
            'executive_summary': True,
            'table_of_contents': True,
            'detailed_analysis': True,
            'recommendations': True,
            'appendices': True,
            'charts': True
        }
    
    def _safe_print(self, message: str):
        """Print message safely, handling encoding issues"""
        try:
            print(message)
        except UnicodeEncodeError:
            # Fallback to ASCII-safe message
            safe_message = message.encode('ascii', errors='ignore').decode('ascii')
            print(safe_message)
    
    def generate_professional_report(self, intelligence_data: Dict[str, Any], 
                                   output_path: str,
                                   template_config: Optional[Dict[str, Any]] = None) -> str:
        """
        Generate a professional Word document report
        
        Args:
            intelligence_data: Application intelligence data
            output_path: Output file path
            template_config: Optional template configuration
            
        Returns:
            Path to generated document
        """
        self._safe_print("[WORD-GEN] Starting professional Word document generation...")
        
        # Import Document here to avoid import errors
        from docx import Document
        
        # Load template or create new document
        if self.template_path and os.path.exists(self.template_path):
            doc = Document(self.template_path)
            self._safe_print(f"[WORD-GEN] Using template: {self.template_path}")
        else:
            doc = Document()
            self._safe_print("[WORD-GEN] Creating new document")
        
        # Apply template configuration
        if template_config:
            self.template_structure.update(template_config)
        
        # Set up document styles
        self._setup_document_styles(doc)
        
        # Generate document sections
        self._generate_cover_page(doc, intelligence_data)
        
        if self.template_structure.get('table_of_contents', True):
            self._generate_table_of_contents(doc)
        
        if self.template_structure.get('executive_summary', True):
            self._generate_executive_summary(doc, intelligence_data)
        
        if self.template_structure.get('detailed_analysis', True):
            self._generate_detailed_analysis(doc, intelligence_data)
        
        if self.template_structure.get('recommendations', True):
            self._generate_recommendations(doc, intelligence_data)
        
        if self.template_structure.get('appendices', True):
            self._generate_appendices(doc, intelligence_data)
        
        # Generate charts if enabled
        if self.template_structure.get('charts', True) and MATPLOTLIB_AVAILABLE:
            self._generate_enhanced_charts(doc, intelligence_data)
        
        # Save document
        doc.save(output_path)
        self._safe_print(f"✅ [WORD-GEN] Professional report saved to: {output_path}")
        
        return output_path
    
    def _setup_document_styles(self, doc):
        """Set up professional document styles"""
        self._safe_print("🎨 [WORD-GEN] Setting up document styles...")
        
        styles = doc.styles
        
        # Title style
        if 'Custom Title' not in [s.name for s in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = 'Segoe UI'
            title_style.font.size = Pt(28)
            title_style.font.bold = True
            title_style.font.color.rgb = self.colors['primary']
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(24)
        
        # Heading styles
        if 'Custom Heading 1' not in [s.name for s in styles]:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_style.font.name = 'Segoe UI'
            h1_style.font.size = Pt(20)
            h1_style.font.bold = True
            h1_style.font.color.rgb = self.colors['primary']
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
        
        if 'Custom Heading 2' not in [s.name for s in styles]:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_style.font.name = 'Segoe UI'
            h2_style.font.size = Pt(16)
            h2_style.font.bold = True
            h2_style.font.color.rgb = self.colors['secondary']
            h2_style.paragraph_format.space_before = Pt(14)
            h2_style.paragraph_format.space_after = Pt(8)
        
        # Body text style
        if 'Custom Body' not in [s.name for s in styles]:
            body_style = styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Segoe UI'
            body_style.font.size = Pt(11)
            body_style.font.color.rgb = self.colors['secondary']
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.line_spacing = 1.15
        
        # Bullet point style
        if 'Custom Bullet' not in [s.name for s in styles]:
            bullet_style = styles.add_style('Custom Bullet', WD_STYLE_TYPE.PARAGRAPH)
            bullet_style.font.name = 'Segoe UI'
            bullet_style.font.size = Pt(11)
            bullet_style.font.color.rgb = self.colors['secondary']
            bullet_style.paragraph_format.space_after = Pt(3)
            bullet_style.paragraph_format.left_indent = Inches(0.25)
    
    def _generate_cover_page(self, doc, intelligence_data: Dict[str, Any]):
        """Generate professional cover page"""
        self._safe_print("📑 [WORD-GEN] Generating cover page...")
        
        # Company logo placeholder
        logo_paragraph = doc.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run("[COMPANY LOGO]")
        logo_run.font.size = Pt(12)
        logo_run.font.color.rgb = self.colors['secondary']
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Main title
        title = doc.add_paragraph()
        title.style = 'Custom Title'
        title_run = title.add_run("Application Intelligence Report")
        
        # Subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.add_run("Comprehensive Analysis and Migration Assessment")
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = self.colors['secondary']
        
        # Add spacing
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Repository information
        repo_info = doc.add_paragraph()
        repo_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        repo_url = intelligence_data.get('repository_url', 'Unknown Repository')
        repo_info.add_run(f"Repository: {repo_url}").font.size = Pt(14)
        
        # Analysis date
        analysis_date = intelligence_data.get('analysis_timestamp', datetime.now().isoformat())
        date_formatted = datetime.fromisoformat(analysis_date.replace('Z', '+00:00')).strftime('%B %d, %Y')
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f"Analysis Date: {date_formatted}").font.size = Pt(12)
        
        # Add spacing
        for _ in range(6):
            doc.add_paragraph()
        
        # Footer information
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run("Generated by Application Intelligence Platform")
        footer_run.font.size = Pt(10)
        footer_run.font.color.rgb = self.colors['secondary']
        footer_run.italic = True
        
        # Page break
        doc.add_page_break()
    
    def _generate_table_of_contents(self, doc):
        """Generate table of contents"""
        self._safe_print("📚 [WORD-GEN] Generating table of contents...")
        
        heading = doc.add_paragraph()
        heading.style = 'Custom Heading 1'
        heading.add_run("Table of Contents")
        
        # TOC entries
        toc_items = [
            ("Executive Summary", "3"),
            ("Application Overview", "4"),
            ("Component Analysis", "5"),
            ("Architecture Assessment", "8"),
            ("Security Analysis", "10"),
            ("Git History Analysis", "12"),
            ("Recommendations", "13"),
            ("Appendices", "15")
        ]
        
        for item, page in toc_items:
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.left_indent = Inches(0.25)
            toc_para.add_run(f"{item}").font.size = Pt(11)
            
            # Add tab leader
            tab_run = toc_para.add_run(f"{'.' * (60 - len(item))} {page}")
            tab_run.font.size = Pt(11)
            tab_run.font.color.rgb = self.colors['secondary']
        
        doc.add_page_break()
    
    def _generate_executive_summary(self, doc, intelligence_data: Dict[str, Any]):
        """Generate executive summary section"""
        self._safe_print("📊 [WORD-GEN] Generating executive summary...")
        
        # Section heading
        heading = doc.add_paragraph()
        heading.style = 'Custom Heading 1'
        heading.add_run("Executive Summary")
        
        # Key metrics table
        self._add_key_metrics_table(doc, intelligence_data)
        
        # Summary text
        summary_data = intelligence_data.get('summary', {})
        
        # Application overview
        overview_heading = doc.add_paragraph()
        overview_heading.style = 'Custom Heading 2'
        overview_heading.add_run("Application Overview")
        
        overview_text = f"""
This report presents a comprehensive analysis of the application repository. The analysis identified {summary_data.get('total_components', 0)} components using {len(summary_data.get('languages', []))} different programming languages. The application demonstrates {'a microservices' if summary_data.get('total_components', 0) > 1 else 'a monolithic'} architecture pattern.
        """
        
        overview_para = doc.add_paragraph()
        overview_para.style = 'Custom Body'
        overview_para.add_run(overview_text.strip())
        
        # Key findings
        findings_heading = doc.add_paragraph()
        findings_heading.style = 'Custom Heading 2'
        findings_heading.add_run("Key Findings")
        
        findings = [
            f"📊 {summary_data.get('total_components', 0)} application components identified",
            f"🔧 {len(summary_data.get('languages', []))} programming languages detected: {', '.join(summary_data.get('languages', []))}",
            f"🐳 {summary_data.get('containerization_status', 0)} components are containerized",
            f"🗃️ {summary_data.get('datasources', 0)} data sources identified",
            f"🔐 {summary_data.get('security_findings', {}).get('vulnerabilities', 0)} security findings require attention"
        ]
        
        for finding in findings:
            finding_para = doc.add_paragraph()
            finding_para.style = 'Custom Bullet'
            finding_para.add_run(f"• {finding}")
        
        doc.add_page_break()
    
    def _add_key_metrics_table(self, doc, intelligence_data: Dict[str, Any]):
        """Add key metrics table"""
        summary_data = intelligence_data.get('summary', {})
        
        # Create table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        # Style header
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = self.colors['primary']
        
        # Data rows
        metrics = [
            ('Total Components', str(summary_data.get('total_components', 0))),
            ('Programming Languages', ', '.join(summary_data.get('languages', []))),
            ('Containerization Status', f"{summary_data.get('containerization_status', 0)} containerized"),
            ('Data Sources', str(summary_data.get('datasources', 0))),
            ('Security Findings', str(summary_data.get('security_findings', {}).get('vulnerabilities', 0))),
            ('Git Commits', str(summary_data.get('git_history', {}).get('total_commits', 0))),
            ('Architecture Style', summary_data.get('architecture_style', {}).get('style', 'Unknown') if isinstance(summary_data.get('architecture_style'), dict) else str(summary_data.get('architecture_style', 'Unknown')))
        ]
        
        for metric, value in metrics:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Add spacing after table
        doc.add_paragraph()
    
    def _generate_detailed_analysis(self, doc, intelligence_data: Dict[str, Any]):
        """Generate detailed analysis section"""
        self._safe_print("🔍 [WORD-GEN] Generating detailed analysis...")
        
        # Section heading
        heading = doc.add_paragraph()
        heading.style = 'Custom Heading 1'
        heading.add_run("Detailed Analysis")
        
        # Component analysis
        self._generate_component_analysis(doc, intelligence_data)
        
        # Architecture analysis
        self._generate_architecture_analysis(doc, intelligence_data)
        
        # Security analysis
        self._generate_security_analysis(doc, intelligence_data)
        
        # Git history analysis
        self._generate_git_analysis(doc, intelligence_data)
    
    def _generate_component_analysis(self, doc, intelligence_data: Dict[str, Any]):
        """Generate component analysis section"""
        components_heading = doc.add_paragraph()
        components_heading.style = 'Custom Heading 2'
        components_heading.add_run("Component Analysis")
        
        components = intelligence_data.get('components', {})
        
        if not components:
            no_components_para = doc.add_paragraph()
            no_components_para.style = 'Custom Body'
            no_components_para.add_run("No components were identified in this analysis.")
            return
        
        # Component summary
        summary_para = doc.add_paragraph()
        summary_para.style = 'Custom Body'
        summary_para.add_run(f"The analysis identified {len(components)} components across the application:")
        
        # Component details table
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # Header
        header_cells = table.rows[0].cells
        headers = ['Component', 'Language', 'Type', 'Packaging']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
            header_cells[i].paragraphs[0].runs[0].font.color.rgb = self.colors['primary']
        
        # Component rows
        for comp_name, comp_data in components.items():
            row_cells = table.add_row().cells
            row_cells[0].text = comp_name
            row_cells[1].text = comp_data.get('language', 'Unknown')
            row_cells[2].text = comp_data.get('type', 'Unknown')
            row_cells[3].text = comp_data.get('packaging', 'Unknown')
        
        doc.add_paragraph()
        
        # Detailed component information
        for comp_name, comp_data in components.items():
            comp_heading = doc.add_paragraph()
            comp_heading.style = 'Custom Heading 2'
            comp_heading.add_run(f"Component: {comp_name}")
            
            details = [
                f"Language: {comp_data.get('language', 'Unknown')}",
                f"Runtime: {comp_data.get('runtime', 'Unknown')}",
                f"Build Tool: {comp_data.get('build_tool', 'Unknown')}",
                f"Packaging: {comp_data.get('packaging', 'Unknown')}"
            ]
            
            # Add exposed ports if available
            if comp_data.get('exposed_ports'):
                details.append(f"Exposed Ports: {', '.join(map(str, comp_data['exposed_ports']))}")
            
            # Add base images if available
            if comp_data.get('base_images'):
                details.append(f"Base Images: {', '.join(comp_data['base_images'])}")
            
            for detail in details:
                detail_para = doc.add_paragraph()
                detail_para.style = 'Custom Bullet'
                detail_para.add_run(f"• {detail}")
            
            # Add notes if available
            if comp_data.get('notes'):
                notes_para = doc.add_paragraph()
                notes_para.style = 'Custom Body'
                notes_para.add_run("Notes:")
                notes_para.runs[0].font.bold = True
                
                for note in comp_data['notes']:
                    note_para = doc.add_paragraph()
                    note_para.style = 'Custom Bullet'
                    note_para.add_run(f"• {note}")
            
            doc.add_paragraph()
    
    def _generate_architecture_analysis(self, doc, intelligence_data: Dict[str, Any]):
        """Generate architecture analysis section"""
        arch_heading = doc.add_paragraph()
        arch_heading.style = 'Custom Heading 2'
        arch_heading.add_run("Architecture Analysis")
        
        arch_assessment = intelligence_data.get('architecture_assessment', {})
        
        if isinstance(arch_assessment, dict):
            style_info = arch_assessment.get('style', {})
            if isinstance(style_info, dict):
                style = style_info.get('value', 'Unknown')
                confidence = style_info.get('confidence', 'Unknown')
                evidence = style_info.get('evidence', [])
                reasoning = style_info.get('reasoning', '')
                
                arch_para = doc.add_paragraph()
                arch_para.style = 'Custom Body'
                arch_para.add_run(f"Architecture Style: {style} (Confidence: {confidence})")
                
                if reasoning:
                    reasoning_para = doc.add_paragraph()
                    reasoning_para.style = 'Custom Body'
                    reasoning_para.add_run(f"Reasoning: {reasoning}")
                
                if evidence:
                    evidence_para = doc.add_paragraph()
                    evidence_para.style = 'Custom Body'
                    evidence_para.add_run("Evidence:")
                    evidence_para.runs[0].font.bold = True
                    
                    for item in evidence:
                        evidence_item = doc.add_paragraph()
                        evidence_item.style = 'Custom Bullet'
                        evidence_item.add_run(f"• {item}")
        
        doc.add_paragraph()
    
    def _generate_security_analysis(self, doc, intelligence_data: Dict[str, Any]):
        """Generate security analysis section"""
        security_heading = doc.add_paragraph()
        security_heading.style = 'Custom Heading 2'
        security_heading.add_run("Security Analysis")
        
        security_data = intelligence_data.get('summary', {}).get('security_findings', {})
        vuln_assessment = intelligence_data.get('vulnerability_assessment', {})
        
        # Security summary
        summary_para = doc.add_paragraph()
        summary_para.style = 'Custom Body'
        vulnerabilities = security_data.get('vulnerabilities', 0)
        
        if isinstance(vuln_assessment, dict):
            findings_count = len(vuln_assessment.get('findings', []))
            
            # Handle both list and integer values for base_image_risks
            base_image_risks_data = vuln_assessment.get('base_image_risks', [])
            if isinstance(base_image_risks_data, list):
                base_image_risks = len(base_image_risks_data)
            else:
                base_image_risks = base_image_risks_data  # Already an integer
            
            summary_text = f"Security analysis identified {findings_count} findings with {base_image_risks} base image risks."
        else:
            summary_text = f"Security analysis identified {vulnerabilities} potential vulnerabilities."
        
        summary_para.add_run(summary_text)
        
        # Detailed findings
        if isinstance(vuln_assessment, dict) and vuln_assessment.get('findings'):
            findings_para = doc.add_paragraph()
            findings_para.style = 'Custom Body'
            findings_para.add_run("Key Security Findings:")
            findings_para.runs[0].font.bold = True
            
            for finding in vuln_assessment['findings'][:5]:  # Top 5 findings
                finding_para = doc.add_paragraph()
                finding_para.style = 'Custom Bullet'
                component = finding.get('component', 'Unknown')
                description = finding.get('description', 'No description')
                severity = finding.get('severity', 'Unknown')
                finding_para.add_run(f"• {component}: {description} (Severity: {severity})")
        
        doc.add_paragraph()
    
    def _generate_git_analysis(self, doc, intelligence_data: Dict[str, Any]):
        """Generate git analysis section"""
        git_heading = doc.add_paragraph()
        git_heading.style = 'Custom Heading 2'
        git_heading.add_run("Git History Analysis")
        
        git_data = intelligence_data.get('git_history', {})
        
        if not git_data:
            no_git_para = doc.add_paragraph()
            no_git_para.style = 'Custom Body'
            no_git_para.add_run("Git history analysis was not available for this repository.")
            return
        
        # Git metrics
        metrics = [
            f"Total Commits: {git_data.get('total_commits', 0)}",
            f"Active Contributors: {git_data.get('active_contributors', 0)}",
            f"Recent Activity: {git_data.get('recent_activity', 'Unknown')}",
            f"Code Stability: {git_data.get('code_stability', 'Unknown')}"
        ]
        
        for metric in metrics:
            metric_para = doc.add_paragraph()
            metric_para.style = 'Custom Bullet'
            metric_para.add_run(f"• {metric}")
        
        doc.add_paragraph()
    
    def _format_recommendation_text(self, rec: str) -> str:
        """Format recommendation text by cleaning up any remaining placeholder tags"""
        # Remove any remaining placeholder tags if they exist
        cleaned_rec = rec
        
        # If the recommendation already has an icon, don't add another
        if any(emoji in rec for emoji in ['🔐', '🐳', '💼', '📈', '🔥', '⚠️', '🚨', 'ℹ️']):
            return cleaned_rec
        
        return cleaned_rec
    
    def _generate_recommendations(self, doc, intelligence_data: Dict[str, Any]):
        """Generate recommendations section"""
        self._safe_print("📝 [WORD-GEN] Generating recommendations...")
        
        heading = doc.add_paragraph()
        heading.style = 'Custom Heading 1'
        heading.add_run("Recommendations")
        
        recommendations = intelligence_data.get('recommendations', [])
        
        if not recommendations:
            no_rec_para = doc.add_paragraph()
            no_rec_para.style = 'Custom Body'
            no_rec_para.add_run("No specific recommendations were generated for this analysis.")
            return
        
        # Priority categorization
        high_priority = []
        medium_priority = []
        low_priority = []
        
        for rec in recommendations:
            formatted_rec = self._format_recommendation_text(rec)
            if any(keyword in rec.lower() for keyword in ['critical', 'security', 'vulnerability', 'urgent']):
                high_priority.append(formatted_rec)
            elif any(keyword in rec.lower() for keyword in ['performance', 'optimization', 'improvement']):
                medium_priority.append(formatted_rec)
            else:
                low_priority.append(formatted_rec)
        
        # High priority recommendations
        if high_priority:
            high_heading = doc.add_paragraph()
            high_heading.style = 'Custom Heading 2'
            high_heading.add_run("🚨 High Priority Recommendations")
            
            for rec in high_priority:
                rec_para = doc.add_paragraph()
                rec_para.style = 'Custom Bullet'
                rec_para.add_run(f"• {rec}")
        
        # Medium priority recommendations
        if medium_priority:
            medium_heading = doc.add_paragraph()
            medium_heading.style = 'Custom Heading 2'
            medium_heading.add_run("⚠️ Medium Priority Recommendations")
            
            for rec in medium_priority:
                rec_para = doc.add_paragraph()
                rec_para.style = 'Custom Bullet'
                rec_para.add_run(f"• {rec}")
        
        # Low priority recommendations
        if low_priority:
            low_heading = doc.add_paragraph()
            low_heading.style = 'Custom Heading 2'
            low_heading.add_run("ℹ️ Low Priority Recommendations")
            
            for rec in low_priority:
                rec_para = doc.add_paragraph()
                rec_para.style = 'Custom Bullet'
                rec_para.add_run(f"• {rec}")
        
        doc.add_page_break()
    
    def _generate_appendices(self, doc, intelligence_data: Dict[str, Any]):
        """Generate appendices section"""
        self._safe_print("📋 [WORD-GEN] Generating appendices...")
        
        heading = doc.add_paragraph()
        heading.style = 'Custom Heading 1'
        heading.add_run("Appendices")
        
        # Appendix A: Technical Details
        appendix_a = doc.add_paragraph()
        appendix_a.style = 'Custom Heading 2'
        appendix_a.add_run("Appendix A: Technical Details")
        
        tech_para = doc.add_paragraph()
        tech_para.style = 'Custom Body'
        tech_para.add_run("This analysis was generated using the Application Intelligence Platform, which performs comprehensive analysis of application repositories including code structure, infrastructure configuration, and security assessment.")
        
        # Appendix B: Methodology
        appendix_b = doc.add_paragraph()
        appendix_b.style = 'Custom Heading 2'
        appendix_b.add_run("Appendix B: Analysis Methodology")
        
        methodology_items = [
            "Component Discovery: Automated scanning of source code and configuration files",
            "Language Detection: Analysis of file extensions, build configurations, and base images",
            "Architecture Assessment: Evaluation of deployment patterns and component relationships",
            "Security Analysis: Scanning for common vulnerabilities and configuration issues",
            "Git History Analysis: Examination of commit patterns and development activity"
        ]
        
        for item in methodology_items:
            method_para = doc.add_paragraph()
            method_para.style = 'Custom Bullet'
            method_para.add_run(f"• {item}")
        
        doc.add_paragraph()
    
    def _generate_enhanced_charts(self, doc, intelligence_data: Dict[str, Any]):
        """Generate enhanced charts and visualizations"""
        self._safe_print("📊 [WORD-GEN] Generating enhanced charts and visualizations...")
        
        if not MATPLOTLIB_AVAILABLE:
            self._safe_print("⚠️  [WORD-GEN] Matplotlib not available, skipping charts")
            return
        
        charts_heading = doc.add_paragraph()
        charts_heading.style = 'Custom Heading 1'
        charts_heading.add_run("Charts and Visualizations")
        
        # Create charts directory
        charts_dir = Path("charts")
        charts_dir.mkdir(exist_ok=True)
        
        try:
            # Import enhanced diagram generator
            from analysis.enhanced_diagram_generator import EnhancedDiagramGenerator
            
            # Generate all enhanced diagrams
            diagram_generator = EnhancedDiagramGenerator(charts_dir)
            generated_diagrams = diagram_generator.generate_all_diagrams(intelligence_data)
            
            # Generate Mermaid and Graphviz diagrams
            try:
                from analysis.mermaid_graphviz_generator import MermaidGraphvizGenerator
                
                mermaid_generator = MermaidGraphvizGenerator(charts_dir)
                mermaid_diagrams = mermaid_generator.generate_all_professional_diagrams(intelligence_data)
                generated_diagrams.extend(mermaid_diagrams)
                
                self._safe_print(f"📊 [WORD-GEN] Added {len(mermaid_diagrams)} Mermaid/Graphviz diagrams")
                
            except ImportError as e:
                self._safe_print(f"⚠️  [WORD-GEN] Mermaid/Graphviz generator not available: {e}")
            except Exception as e:
                self._safe_print(f"⚠️  [WORD-GEN] Error generating Mermaid/Graphviz diagrams: {e}")
            
            # Insert diagrams into document
            diagram_titles = {
                # Enhanced matplotlib diagrams
                'language_distribution.png': 'Programming Language Distribution',
                'component_status.png': 'Component Status Overview',
                'architecture_diagram.png': 'Application Architecture',
                'security_radar.png': 'Security Assessment Radar',
                'vulnerability_timeline.png': 'Vulnerability Analysis',
                'component_network.png': 'Component Network Topology',
                'technology_stack.png': 'Technology Stack',
                'development_activity.png': 'Development Activity Heatmap',
                'migration_readiness.png': 'Migration Readiness Assessment',
                'risk_assessment_matrix.png': 'Risk Assessment Matrix',
                'cicd_pipeline.png': 'CI/CD Pipeline Overview',
                'user_flow_diagram.png': 'User Flow Diagram',
                'high_level_architecture.png': 'High-Level Architecture Overview',
                'system_overview.png': 'System Overview Diagram',
                
                # Mermaid diagrams
                'architecture_flowchart.png': 'Architecture Flowchart (Mermaid)',
                'security_flow_diagram.png': 'Security Flow Diagram (Mermaid)',
                'cicd_pipeline_mermaid.png': 'CI/CD Pipeline Flow (Mermaid)',
                'deployment_architecture.png': 'Deployment Architecture (Mermaid)',
                'risk_assessment_flow.png': 'Risk Assessment Flow (Mermaid)',
                'user_journey_map.png': 'User Journey Map (Mermaid)',
                'business_flow_diagram.png': 'Business Flow Diagram (Mermaid)',
                
                # Graphviz diagrams
                'component_relationship_graph.png': 'Component Relationships (Graphviz)',
                'data_flow_diagram.png': 'Data Flow Diagram (Graphviz)',
                'migration_strategy_diagram.png': 'Migration Strategy (Graphviz)'
            }
            
            # Import diagram summary analyzer
            try:
                from analysis.diagram_summary_analyzer import DiagramSummaryAnalyzer
                summary_analyzer = DiagramSummaryAnalyzer()
                use_summaries = True
                self._safe_print("📊 [WORD-GEN] Diagram summary analyzer loaded")
            except ImportError:
                use_summaries = False
                self._safe_print("⚠️  [WORD-GEN] Diagram summary analyzer not available")
            
            for diagram_path in generated_diagrams:
                try:
                    diagram_file = Path(diagram_path).name
                    title = diagram_titles.get(diagram_file, diagram_file.replace('.png', '').replace('_', ' ').title())
                    
                    # Add diagram heading
                    diagram_para = doc.add_paragraph()
                    diagram_para.style = 'Custom Heading 2'
                    diagram_para.add_run(title)
                    
                    # Insert diagram
                    doc.add_picture(str(diagram_path), width=Inches(6))
                    
                    # Add intelligent summary if available
                    if use_summaries:
                        try:
                            summary = summary_analyzer.generate_diagram_summary(diagram_file, intelligence_data)
                            self._add_diagram_summary(doc, summary)
                        except Exception as e:
                            self._safe_print(f"⚠️  [WORD-GEN] Error generating summary for {diagram_file}: {e}")
                    
                    doc.add_paragraph()
                    
                except Exception as e:
                    self._safe_print(f"⚠️  [WORD-GEN] Error inserting diagram {diagram_path}: {e}")
                    
        except ImportError:
            self._safe_print("⚠️  [WORD-GEN] Enhanced diagram generator not available, using basic charts")
            # Fallback to basic charts
            self._generate_basic_charts(doc, intelligence_data, charts_dir)
        except Exception as e:
            self._safe_print(f"⚠️  [WORD-GEN] Error with enhanced diagrams: {e}")
            # Fallback to basic charts
            self._generate_basic_charts(doc, intelligence_data, charts_dir)
    
    def _generate_basic_charts(self, doc, intelligence_data: Dict[str, Any], charts_dir: Path):
        """Generate basic charts (fallback)"""
        # Language distribution chart
        self._create_language_chart(intelligence_data, charts_dir)
        
        # Component status chart
        self._create_component_chart(intelligence_data, charts_dir)
        
        # Insert charts into document
        try:
            lang_chart_path = charts_dir / "language_distribution.png"
            if lang_chart_path.exists():
                lang_chart_para = doc.add_paragraph()
                lang_chart_para.style = 'Custom Heading 2'
                lang_chart_para.add_run("Language Distribution")
                
                doc.add_picture(str(lang_chart_path), width=Inches(5))
                doc.add_paragraph()
            
            comp_chart_path = charts_dir / "component_status.png"
            if comp_chart_path.exists():
                comp_chart_para = doc.add_paragraph()
                comp_chart_para.style = 'Custom Heading 2'
                comp_chart_para.add_run("Component Status")
                
                doc.add_picture(str(comp_chart_path), width=Inches(5))
                doc.add_paragraph()
        except Exception as e:
            self._safe_print(f"⚠️  [WORD-GEN] Error inserting basic charts: {e}")
    
    def _add_diagram_summary(self, doc, summary: Dict[str, str]):
        """Add intelligent summary for a diagram with professional formatting"""
        try:
            # Add summary section with subtle styling
            summary_para = doc.add_paragraph()
            summary_para.style = 'Normal'
            summary_para.space_before = Pt(6)
            summary_para.space_after = Pt(12)
            
            # Add summary sections with color-coded headers
            sections = [
                ('Context', summary.get('context', ''), '#2c3e50'),
                ('Key Insights', summary.get('key_insights', ''), '#2980b9'),
                ('Business Impact', summary.get('business_impact', ''), '#e74c3c'),
                ('Recommendations', summary.get('recommendations', ''), '#27ae60'),
                ('Technical Details', summary.get('technical_details', ''), '#7f8c8d')
            ]
            
            for section_title, content, color in sections:
                if content:
                    # Section header
                    header_para = doc.add_paragraph()
                    header_run = header_para.add_run(f"📋 {section_title}: ")
                    header_run.font.bold = True
                    header_run.font.size = Pt(10)
                    
                    # Section content
                    content_run = header_para.add_run(content)
                    content_run.font.size = Pt(9)
                    
                    # Add spacing
                    header_para.space_after = Pt(6)
            
            # Add separator
            doc.add_paragraph("─" * 80).style = 'Normal'
            
        except Exception as e:
            self._safe_print(f"⚠️  [WORD-GEN] Error adding diagram summary: {e}")
            # Add basic text fallback
            try:
                fallback_para = doc.add_paragraph()
                fallback_para.add_run("Analysis insights available for this diagram.")
                fallback_para.space_after = Pt(12)
            except:
                pass
    
    def _create_language_chart(self, intelligence_data: Dict[str, Any], charts_dir: Path):
        """Create language distribution chart"""
        try:
            languages = intelligence_data.get('summary', {}).get('languages', [])
            
            if not languages:
                return
            
            # Count languages
            lang_counts = {}
            for lang in languages:
                lang_counts[lang] = lang_counts.get(lang, 0) + 1
            
            # Create pie chart
            fig, ax = plt.subplots(figsize=(8, 6))
            colors = ['#007bd3', '#48b0f1', '#78c5f0', '#a8dbef', '#d8f1fe']
            
            wedges, texts, autotexts = ax.pie(
                lang_counts.values(),
                labels=lang_counts.keys(),
                autopct='%1.1f%%',
                colors=colors[:len(lang_counts)],
                startangle=90
            )
            
            ax.set_title('Programming Language Distribution', fontsize=14, fontweight='bold')
            plt.tight_layout()
            plt.savefig(charts_dir / "language_distribution.png", dpi=300, bbox_inches='tight')
            plt.close()
            
        except Exception as e:
            self._safe_print(f"⚠️  [WORD-GEN] Error creating language chart: {e}")
    
    def _create_component_chart(self, intelligence_data: Dict[str, Any], charts_dir: Path):
        """Create component status chart"""
        try:
            summary = intelligence_data.get('summary', {})
            
            # Data for chart
            total_components = summary.get('total_components', 0)
            containerized = summary.get('containerization_status', 0)
            with_security_issues = summary.get('security_findings', {}).get('vulnerabilities', 0)
            
            categories = ['Total Components', 'Containerized', 'Security Issues']
            values = [total_components, containerized, with_security_issues]
            
            # Create bar chart
            fig, ax = plt.subplots(figsize=(10, 6))
            colors = ['#007bd3', '#16a085', '#e74c3c']
            
            bars = ax.bar(categories, values, color=colors)
            
            # Add value labels on bars
            for bar, value in zip(bars, values):
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{value}', ha='center', va='bottom', fontweight='bold')
            
            ax.set_title('Component Status Overview', fontsize=14, fontweight='bold')
            ax.set_ylabel('Count')
            
            plt.tight_layout()
            plt.savefig(charts_dir / "component_status.png", dpi=300, bbox_inches='tight')
            plt.close()
            
        except Exception as e:
            self._safe_print(f"⚠️  [WORD-GEN] Error creating component chart: {e}")

def create_custom_template(template_path: str, config: Dict[str, Any]) -> str:
    """
    Create a custom Word template with specified configuration
    
    Args:
        template_path: Path to save the template
        config: Template configuration
        
    Returns:
        Path to created template
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is required. Install with: pip install python-docx")
    
    print(f"[TEMPLATE] Creating custom template: {template_path}")
    
    from docx import Document
    doc = Document()
    
    # Add template placeholder content
    title = doc.add_paragraph()
    title.add_run("APPLICATION INTELLIGENCE REPORT").font.size = Pt(24)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph("This is a custom template for Application Intelligence Reports.")
    doc.add_paragraph()
    doc.add_paragraph("Template Configuration:")
    
    for key, value in config.items():
        doc.add_paragraph(f"• {key}: {value}")
    
    doc.add_paragraph()
    doc.add_paragraph("Content will be automatically generated here...")
    
    doc.save(template_path)
    print(f"[TEMPLATE] Template created: {template_path}")
    
    return template_path

# Example usage and testing
if __name__ == "__main__":
    # Test the professional Word generator
    if DOCX_AVAILABLE:
        generator = ProfessionalWordGenerator()
        
        # Sample data for testing
        sample_data = {
            'repository_url': 'https://github.com/example/test-app',
            'analysis_timestamp': datetime.now().isoformat(),
            'summary': {
                'total_components': 3,
                'languages': ['python', 'javascript', 'java'],
                'containerization_status': 2,
                'datasources': 1,
                'security_findings': {'vulnerabilities': 3}
            },
            'components': {
                'web-app': {
                    'language': 'python',
                    'type': 'web',
                    'packaging': 'docker',
                    'runtime': 'python3.9',
                    'build_tool': 'pip',
                    'exposed_ports': [8080],
                    'base_images': ['python:3.9-slim']
                }
            },
            'recommendations': [
                'Update base images to latest versions',
                'Implement security scanning in CI/CD pipeline',
                'Add monitoring and logging capabilities'
            ]
        }
        
        # Generate report
        output_path = "test_professional_report.docx"
        generator.generate_professional_report(sample_data, output_path)
        print(f"Test report generated: {output_path}")
    else:
        print("python-docx not available for testing")