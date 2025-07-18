"""
Professional Word Document Generator for Sprint 2 AS-IS Analysis
Creates business-ready Word documents from Sprint 2 analysis results
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from src.core.models import (
    RepositoryAnalysis, ComponentInfo, ServiceDependency, 
    ServiceCriticality, SecurityFindings
)


class Sprint2WordGenerator:
    """Generates professional Word documents for Sprint 2 AS-IS analysis"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word document generation. Install with: pip install python-docx")
        
        self.doc = None
        self.styles = {}
        
    def generate_as_is_word_document(self, analysis: RepositoryAnalysis, output_path: str) -> str:
        """Generate professional Word document from Sprint 2 analysis"""
        
        # Create new document
        self.doc = Document()
        
        # Setup document styles
        self._setup_document_styles()
        
        # Add document properties
        self._add_document_properties(analysis)
        
        # Generate document content
        self._generate_title_page(analysis)
        self._generate_executive_summary(analysis)
        self._generate_table_of_contents()
        self._generate_system_overview(analysis)
        self._generate_component_analysis(analysis)
        self._generate_dependency_analysis(analysis)
        self._generate_criticality_analysis(analysis)
        self._generate_security_analysis(analysis)
        self._generate_architecture_insights(analysis)
        self._generate_migration_recommendations(analysis)
        self._generate_appendices(analysis)
        
        # Save document
        self.doc.save(output_path)
        
        return output_path
    
    def _setup_document_styles(self):
        """Setup professional document styles"""
        
        # Document title style
        title_style = self.doc.styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(18)
        
        # Heading 1 style
        h1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(20)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.page_break_before = True
        
        # Heading 2 style
        h2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(47, 84, 150)  # Medium blue
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Heading 3 style
        h3_style = self.doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(68, 114, 196)  # Light blue
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Normal text style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        
        # Highlighted text style
        highlight_style = self.doc.styles.add_style('HighlightedText', WD_STYLE_TYPE.PARAGRAPH)
        highlight_style.font.name = 'Calibri'
        highlight_style.font.size = Pt(11)
        highlight_style.font.bold = True
        highlight_style.font.color.rgb = RGBColor(0, 102, 0)  # Dark green
        
        # Store styles for reference
        self.styles = {
            'title': title_style,
            'h1': h1_style,
            'h2': h2_style,
            'h3': h3_style,
            'normal': normal_style,
            'highlight': highlight_style
        }
    
    def _add_document_properties(self, analysis: RepositoryAnalysis):
        """Add document properties and metadata"""
        
        core_props = self.doc.core_properties
        repo_name = getattr(analysis, 'repo_name', None)
        if not repo_name and hasattr(analysis, 'repository_url'):
            repo_name = analysis.repository_url.rstrip('/').split('/')[-1]
        if not repo_name:
            repo_name = 'Unknown Repository'
            
        core_props.title = f"AS-IS System Analysis - {repo_name}"
        core_props.author = "Migration Analysis Team"
        core_props.subject = "System Analysis and Migration Assessment"
        core_props.comments = f"AS-IS analysis for {repo_name} migration planning"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
    
    def _generate_title_page(self, analysis: RepositoryAnalysis):
        """Generate professional title page"""
        
        # Title
        title = self.doc.add_paragraph("AS-IS System Analysis Report", style=self.styles['title'])
        
        # Subtitle
        subtitle = self.doc.add_paragraph("Migration Assessment and Planning Document")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.italic = True
        subtitle.runs[0].font.color.rgb = RGBColor(89, 89, 89)
        
        # Add spacing
        self.doc.add_paragraph()
        self.doc.add_paragraph()
        
        # Repository information
        repo_name = getattr(analysis, 'repo_name', None)
        if not repo_name and hasattr(analysis, 'repository_url'):
            repo_name = analysis.repository_url.rstrip('/').split('/')[-1]
        if not repo_name:
            repo_name = 'Unknown Repository'
        
        info_table = self.doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Repository
        info_table.cell(0, 0).text = "Repository"
        info_table.cell(0, 1).text = repo_name
        info_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        
        # Analysis Date
        info_table.cell(1, 0).text = "Analysis Date"
        info_table.cell(1, 1).text = analysis.analysis_date.strftime('%B %d, %Y')
        info_table.cell(1, 0).paragraphs[0].runs[0].font.bold = True
        
        # Components
        info_table.cell(2, 0).text = "Components Analyzed"
        info_table.cell(2, 1).text = str(len(analysis.components))
        info_table.cell(2, 0).paragraphs[0].runs[0].font.bold = True
        
        # Dependencies
        info_table.cell(3, 0).text = "Dependencies Identified"
        info_table.cell(3, 1).text = str(len(analysis.dependencies))
        info_table.cell(3, 0).paragraphs[0].runs[0].font.bold = True
        
        # Add page break
        self.doc.add_page_break()
    
    def _generate_executive_summary(self, analysis: RepositoryAnalysis):
        """Generate executive summary section"""
        
        self.doc.add_heading("Executive Summary", level=1).style = self.styles['h1']
        
        # System metrics
        total_files = sum(comp.files_count for comp in analysis.components)
        total_loc = sum(comp.lines_of_code for comp in analysis.components)
        total_endpoints = sum(comp.api_endpoints_count for comp in analysis.components)
        
        # Create metrics table
        metrics_table = self.doc.add_table(rows=5, cols=2)
        metrics_table.style = 'Table Grid'
        
        metrics_data = [
            ("Total Components", str(len(analysis.components))),
            ("Total Files", f"{total_files:,}"),
            ("Total Lines of Code", f"{total_loc:,}"),
            ("Total API Endpoints", str(total_endpoints)),
            ("External Dependencies", str(len([d for d in analysis.dependencies if d.dependency_type == 'http'])))
        ]
        
        for i, (label, value) in enumerate(metrics_data):
            metrics_table.cell(i, 0).text = label
            metrics_table.cell(i, 1).text = value
            metrics_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
        
        # Criticality overview
        if analysis.criticality_assessments:
            critical_count = len([c for c in analysis.criticality_assessments if c.business_criticality == 'critical'])
            high_count = len([c for c in analysis.criticality_assessments if c.business_criticality == 'high'])
            
            criticality_para = self.doc.add_paragraph()
            criticality_para.add_run("Criticality Assessment: ").font.bold = True
            criticality_para.add_run(f"This analysis identified {critical_count} critical components and {high_count} high-priority components that require special attention during migration planning.")
        
        # Key findings
        if analysis.architecture_insights:
            self.doc.add_paragraph()
            key_findings = self.doc.add_paragraph()
            key_findings.add_run("Key Findings: ").font.bold = True
            key_findings.add_run(analysis.architecture_insights[0] if analysis.architecture_insights else "Analysis completed successfully.")
    
    def _generate_table_of_contents(self):
        """Generate table of contents placeholder"""
        self.doc.add_heading("Table of Contents", level=1).style = self.styles['h1']
        self.doc.add_paragraph("1. System Overview")
        self.doc.add_paragraph("2. Component Analysis")
        self.doc.add_paragraph("3. Dependency Analysis")
        self.doc.add_paragraph("4. Criticality Assessment")
        self.doc.add_paragraph("5. Security Analysis")
        self.doc.add_paragraph("6. Architecture Insights")
        self.doc.add_paragraph("7. Migration Recommendations")
        self.doc.add_paragraph("8. Appendices")
        
        self.doc.add_page_break()
    
    def _generate_system_overview(self, analysis: RepositoryAnalysis):
        """Generate system overview section"""
        
        self.doc.add_heading("System Overview", level=1).style = self.styles['h1']
        
        # Language distribution
        language_dist = {}
        for comp in analysis.components:
            if comp.language not in language_dist:
                language_dist[comp.language] = 0
            language_dist[comp.language] += 1
        
        self.doc.add_heading("Technology Stack", level=2).style = self.styles['h2']
        
        lang_table = self.doc.add_table(rows=len(language_dist) + 1, cols=2)
        lang_table.style = 'Table Grid'
        
        # Header
        lang_table.cell(0, 0).text = "Language/Framework"
        lang_table.cell(0, 1).text = "Components"
        lang_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        lang_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        
        # Data
        for i, (language, count) in enumerate(sorted(language_dist.items()), 1):
            lang_table.cell(i, 0).text = language.title()
            lang_table.cell(i, 1).text = str(count)
        
        self.doc.add_paragraph()
        
        # Architecture patterns
        self.doc.add_heading("Architecture Patterns", level=2).style = self.styles['h2']
        patterns = self._detect_architecture_patterns(analysis)
        for pattern in patterns:
            self.doc.add_paragraph(pattern, style='List Bullet')
    
    def _generate_component_analysis(self, analysis: RepositoryAnalysis):
        """Generate detailed component analysis"""
        
        self.doc.add_heading("Component Analysis", level=1).style = self.styles['h1']
        
        # Sort components by criticality
        sorted_components = sorted(
            analysis.components,
            key=lambda c: self._get_component_criticality_score(c, analysis.criticality_assessments),
            reverse=True
        )
        
        for component in sorted_components:
            self.doc.add_heading(component.name, level=2).style = self.styles['h2']
            
            # Component metrics table
            comp_table = self.doc.add_table(rows=6, cols=2)
            comp_table.style = 'Table Grid'
            
            comp_data = [
                ("Language", component.language),
                ("Files", str(component.files_count)),
                ("Lines of Code", f"{component.lines_of_code:,}"),
                ("API Endpoints", str(component.api_endpoints_count)),
                ("Database Operations", str(component.database_operations_count)),
                ("External HTTP Calls", str(component.http_calls_count))
            ]
            
            for i, (label, value) in enumerate(comp_data):
                comp_table.cell(i, 0).text = label
                comp_table.cell(i, 1).text = value
                comp_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
            
            self.doc.add_paragraph()
            
            # Criticality information
            criticality = self._get_component_criticality(component.name, analysis.criticality_assessments)
            if criticality:
                self.doc.add_heading("Criticality Assessment", level=3).style = self.styles['h3']
                
                crit_table = self.doc.add_table(rows=5, cols=2)
                crit_table.style = 'Table Grid'
                
                crit_data = [
                    ("Business Criticality", criticality.business_criticality),
                    ("Technical Complexity", criticality.technical_complexity),
                    ("User Impact", criticality.user_impact),
                    ("Data Sensitivity", criticality.data_sensitivity),
                    ("Risk Score", f"{criticality.score:.2f}")
                ]
                
                for i, (label, value) in enumerate(crit_data):
                    crit_table.cell(i, 0).text = label
                    crit_table.cell(i, 1).text = str(value)
                    crit_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
                
                self.doc.add_paragraph()
                
                reasoning_para = self.doc.add_paragraph()
                reasoning_para.add_run("Assessment Reasoning: ").font.bold = True
                reasoning_para.add_run(criticality.reasoning)
                self.doc.add_paragraph()
    
    def _generate_dependency_analysis(self, analysis: RepositoryAnalysis):
        """Generate dependency analysis section"""
        
        self.doc.add_heading("Dependency Analysis", level=1).style = self.styles['h1']
        
        # Dependency types summary
        dep_types = {}
        for dep in analysis.dependencies:
            dep_types[dep.dependency_type] = dep_types.get(dep.dependency_type, 0) + 1
        
        self.doc.add_heading("Dependency Types", level=2).style = self.styles['h2']
        
        dep_table = self.doc.add_table(rows=len(dep_types) + 1, cols=2)
        dep_table.style = 'Table Grid'
        
        # Header
        dep_table.cell(0, 0).text = "Dependency Type"
        dep_table.cell(0, 1).text = "Count"
        dep_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        dep_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        
        # Data
        for i, (dep_type, count) in enumerate(sorted(dep_types.items()), 1):
            dep_table.cell(i, 0).text = dep_type.title()
            dep_table.cell(i, 1).text = str(count)
        
        self.doc.add_paragraph()
        
        # Top dependencies
        self.doc.add_heading("Key Dependencies", level=2).style = self.styles['h2']
        
        # Show top 10 dependencies
        top_deps = sorted(analysis.dependencies, key=lambda d: d.dependency_type == 'http', reverse=True)[:10]
        
        if top_deps:
            dep_detail_table = self.doc.add_table(rows=len(top_deps) + 1, cols=4)
            dep_detail_table.style = 'Table Grid'
            
            # Header
            headers = ["Source", "Target", "Type", "Endpoint"]
            for i, header in enumerate(headers):
                dep_detail_table.cell(0, i).text = header
                dep_detail_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, dep in enumerate(top_deps, 1):
                dep_detail_table.cell(i, 0).text = dep.source_component
                dep_detail_table.cell(i, 1).text = dep.target_component
                dep_detail_table.cell(i, 2).text = dep.dependency_type
                dep_detail_table.cell(i, 3).text = dep.endpoint or "N/A"
    
    def _generate_criticality_analysis(self, analysis: RepositoryAnalysis):
        """Generate criticality analysis section"""
        
        self.doc.add_heading("Criticality Assessment", level=1).style = self.styles['h1']
        
        if not analysis.criticality_assessments:
            self.doc.add_paragraph("No criticality assessments available.")
            return
        
        # Criticality distribution
        distribution = self._get_criticality_distribution(analysis.criticality_assessments)
        
        self.doc.add_heading("Criticality Distribution", level=2).style = self.styles['h2']
        
        crit_table = self.doc.add_table(rows=5, cols=2)
        crit_table.style = 'Table Grid'
        
        crit_data = [
            ("Critical", str(distribution['critical'])),
            ("High", str(distribution['high'])),
            ("Medium", str(distribution['medium'])),
            ("Low", str(distribution['low'])),
            ("Total", str(len(analysis.criticality_assessments)))
        ]
        
        for i, (level, count) in enumerate(crit_data):
            crit_table.cell(i, 0).text = level
            crit_table.cell(i, 1).text = count
            crit_table.cell(i, 0).paragraphs[0].runs[0].font.bold = True
        
        self.doc.add_paragraph()
        
        # High-risk components
        high_risk = [c for c in analysis.criticality_assessments if c.business_criticality in ['critical', 'high']]
        if high_risk:
            self.doc.add_heading("High-Risk Components", level=2).style = self.styles['h2']
            
            risk_table = self.doc.add_table(rows=len(high_risk) + 1, cols=4)
            risk_table.style = 'Table Grid'
            
            # Header
            risk_headers = ["Component", "Criticality", "Complexity", "Risk Score"]
            for i, header in enumerate(risk_headers):
                risk_table.cell(0, i).text = header
                risk_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, crit in enumerate(high_risk, 1):
                risk_table.cell(i, 0).text = crit.component_name
                risk_table.cell(i, 1).text = crit.business_criticality
                risk_table.cell(i, 2).text = crit.technical_complexity
                risk_table.cell(i, 3).text = f"{crit.score:.2f}"
    
    def _generate_security_analysis(self, analysis: RepositoryAnalysis):
        """Generate security analysis section"""
        
        self.doc.add_heading("Security Analysis", level=1).style = self.styles['h1']
        
        if not analysis.security_findings:
            self.doc.add_paragraph("No security findings available.")
            return
        
        # Security summary
        total_issues = len(analysis.security_findings)
        critical_issues = len([f for f in analysis.security_findings if f.severity == 'critical'])
        high_issues = len([f for f in analysis.security_findings if f.severity == 'high'])
        
        security_summary = self.doc.add_paragraph()
        security_summary.add_run("Security Summary: ").font.bold = True
        security_summary.add_run(f"Analysis identified {total_issues} security issues, including {critical_issues} critical and {high_issues} high-severity findings.")
        
        self.doc.add_paragraph()
        
        # Security findings table
        if analysis.security_findings:
            self.doc.add_heading("Security Findings", level=2).style = self.styles['h2']
            
            sec_table = self.doc.add_table(rows=len(analysis.security_findings) + 1, cols=4)
            sec_table.style = 'Table Grid'
            
            # Header
            sec_headers = ["Component", "Issue Type", "Severity", "Description"]
            for i, header in enumerate(sec_headers):
                sec_table.cell(0, i).text = header
                sec_table.cell(0, i).paragraphs[0].runs[0].font.bold = True
            
            # Data
            for i, finding in enumerate(analysis.security_findings, 1):
                sec_table.cell(i, 0).text = finding.component_name
                sec_table.cell(i, 1).text = finding.issue_type
                sec_table.cell(i, 2).text = finding.severity
                sec_table.cell(i, 3).text = finding.description[:100] + "..." if len(finding.description) > 100 else finding.description
    
    def _generate_architecture_insights(self, analysis: RepositoryAnalysis):
        """Generate architecture insights section"""
        
        self.doc.add_heading("Architecture Insights", level=1).style = self.styles['h1']
        
        if not analysis.architecture_insights:
            self.doc.add_paragraph("No architecture insights available.")
            return
        
        for i, insight in enumerate(analysis.architecture_insights, 1):
            insight_para = self.doc.add_paragraph()
            insight_para.add_run(f"Insight {i}: ").font.bold = True
            insight_para.add_run(insight)
            self.doc.add_paragraph()
    
    def _generate_migration_recommendations(self, analysis: RepositoryAnalysis):
        """Generate migration recommendations section"""
        
        self.doc.add_heading("Migration Recommendations", level=1).style = self.styles['h1']
        
        if not analysis.migration_recommendations:
            self.doc.add_paragraph("No migration recommendations available.")
            return
        
        for i, recommendation in enumerate(analysis.migration_recommendations, 1):
            rec_para = self.doc.add_paragraph()
            rec_para.add_run(f"Recommendation {i}: ").font.bold = True
            rec_para.add_run(recommendation)
            self.doc.add_paragraph()
    
    def _generate_appendices(self, analysis: RepositoryAnalysis):
        """Generate appendices section"""
        
        self.doc.add_heading("Appendices", level=1).style = self.styles['h1']
        
        # Appendix A: Glossary
        self.doc.add_heading("Appendix A: Glossary", level=2).style = self.styles['h2']
        
        glossary_data = [
            ("API", "Application Programming Interface"),
            ("LOC", "Lines of Code"),
            ("HTTP", "Hypertext Transfer Protocol"),
            ("REST", "Representational State Transfer"),
            ("JSON", "JavaScript Object Notation"),
            ("SQL", "Structured Query Language"),
            ("NoSQL", "Not Only SQL"),
            ("CI/CD", "Continuous Integration/Continuous Deployment")
        ]
        
        glossary_table = self.doc.add_table(rows=len(glossary_data) + 1, cols=2)
        glossary_table.style = 'Table Grid'
        
        # Header
        glossary_table.cell(0, 0).text = "Term"
        glossary_table.cell(0, 1).text = "Definition"
        glossary_table.cell(0, 0).paragraphs[0].runs[0].font.bold = True
        glossary_table.cell(0, 1).paragraphs[0].runs[0].font.bold = True
        
        # Data
        for i, (term, definition) in enumerate(glossary_data, 1):
            glossary_table.cell(i, 0).text = term
            glossary_table.cell(i, 1).text = definition
    
    def _detect_architecture_patterns(self, analysis: RepositoryAnalysis) -> List[str]:
        """Detect architecture patterns from analysis"""
        patterns = []
        
        # Microservices pattern
        if len(analysis.components) > 3:
            patterns.append("Microservices Architecture - Multiple independent components identified")
        
        # API-first pattern
        total_endpoints = sum(comp.api_endpoints_count for comp in analysis.components)
        if total_endpoints > 5:
            patterns.append("API-First Design - Significant number of API endpoints detected")
        
        # Database-centric pattern
        total_db_ops = sum(comp.database_operations_count for comp in analysis.components)
        if total_db_ops > 10:
            patterns.append("Database-Centric - High number of database operations")
        
        # Event-driven pattern
        http_calls = sum(comp.http_calls_count for comp in analysis.components)
        if http_calls > 5:
            patterns.append("Event-Driven - Multiple HTTP service calls detected")
        
        if not patterns:
            patterns.append("Monolithic Architecture - Single or few components with tight coupling")
        
        return patterns
    
    def _get_component_criticality_score(self, component: ComponentInfo, assessments: List[ServiceCriticality]) -> float:
        """Get criticality score for component"""
        for assessment in assessments:
            if assessment.component_name == component.name:
                return assessment.score
        return 0.0
    
    def _get_component_criticality(self, component_name: str, assessments: List[ServiceCriticality]) -> Optional[ServiceCriticality]:
        """Get criticality assessment for component"""
        for assessment in assessments:
            if assessment.component_name == component_name:
                return assessment
        return None
    
    def _get_criticality_distribution(self, assessments: List[ServiceCriticality]) -> Dict[str, int]:
        """Get distribution of criticality levels"""
        distribution = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
        for assessment in assessments:
            level = assessment.business_criticality.lower()
            if level in distribution:
                distribution[level] += 1
        return distribution


def generate_sprint2_word_document(analysis: RepositoryAnalysis, output_dir: str) -> str:
    """Convenience function to generate Sprint 2 Word document"""
    
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Generate filename
    repo_name = getattr(analysis, 'repo_name', None)
    if not repo_name and hasattr(analysis, 'repository_url'):
        # Extract repo name from URL
        url_parts = analysis.repository_url.rstrip('/').split('/')
        repo_name = url_parts[-1] if url_parts else 'Unknown'
        
        # If it's a file path, use the directory name
        if repo_name.startswith('file://'):
            path_parts = repo_name.split('/')
            repo_name = path_parts[-1] if len(path_parts) > 1 else 'Local_Project'
    
    if not repo_name:
        repo_name = 'Unknown_Repository'
    
    # Clean repo name for filename and limit length
    repo_name = re.sub(r'[<>:"/\\|?*]', '_', repo_name)
    repo_name = repo_name[:30]  # Limit length
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"ASIS_Analysis_{repo_name}_{timestamp}.docx"
    output_path = os.path.join(output_dir, filename)
    
    # Generate document
    generator = Sprint2WordGenerator()
    result_path = generator.generate_as_is_word_document(analysis, output_path)
    
    return result_path


if __name__ == "__main__":
    print("Sprint 2 Word Document Generator")
    print("This module provides professional Word document generation for Sprint 2 AS-IS analysis.") 