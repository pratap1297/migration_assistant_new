"""
Markdown to Word Document Converter
Converts markdown content to professional Word documents with proper formatting
"""

import os
import re
from datetime import datetime
from typing import Dict, List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class MarkdownToWordConverter:
    """Converts markdown content to professional Word documents"""
    
    def __init__(self):
        if not HAS_DOCX:
            raise ImportError("python-docx is required for Word document generation. Install with: pip install python-docx")
        
        self.doc = None
        self.styles = {}
        
    def convert_markdown_to_word(self, markdown_content: str, output_path: str, title: str = "Document") -> str:
        """Convert markdown content to Word document"""
        
        # Create new document
        self.doc = Document()
        
        # Setup document styles
        self._setup_document_styles()
        
        # Add document properties
        self._add_document_properties(title)
        
        # Parse markdown and convert to Word
        self._convert_markdown_content(markdown_content)
        
        # Save document
        self.doc.save(output_path)
        
        return output_path
    
    def _setup_document_styles(self):
        """Setup professional document styles"""
        
        # Document title style
        title_style = self.doc.styles.add_style('DocumentTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Calibri'
        title_style.font.size = Pt(28)
        title_style.font.bold = True
        title_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(18)
        
        # Heading 1 style
        h1_style = self.doc.styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Calibri'
        h1_style.font.size = Pt(20)
        h1_style.font.bold = True
        h1_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        h1_style.paragraph_format.space_before = Pt(18)
        h1_style.paragraph_format.space_after = Pt(12)
        h1_style.paragraph_format.page_break_before = True
        
        # Heading 2 style
        h2_style = self.doc.styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Calibri'
        h2_style.font.size = Pt(16)
        h2_style.font.bold = True
        h2_style.font.color.rgb = RGBColor(47, 84, 150)  # Medium blue
        h2_style.paragraph_format.space_before = Pt(12)
        h2_style.paragraph_format.space_after = Pt(6)
        
        # Heading 3 style
        h3_style = self.doc.styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
        h3_style.font.name = 'Calibri'
        h3_style.font.size = Pt(14)
        h3_style.font.bold = True
        h3_style.font.color.rgb = RGBColor(68, 114, 196)  # Light blue
        h3_style.paragraph_format.space_before = Pt(10)
        h3_style.paragraph_format.space_after = Pt(4)
        
        # Normal text style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Calibri'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15
        
        # Code style
        code_style = self.doc.styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = 'Consolas'
        code_style.font.size = Pt(10)
        code_style.paragraph_format.left_indent = Inches(0.25)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        
        # Store styles for reference
        self.styles = {
            'title': title_style,
            'h1': h1_style,
            'h2': h2_style,
            'h3': h3_style,
            'normal': normal_style,
            'code': code_style
        }
    
    def _add_document_properties(self, title: str):
        """Add document properties and metadata"""
        
        core_props = self.doc.core_properties
        core_props.title = title
        core_props.author = "Migration Analysis Team"
        core_props.subject = "System Analysis and Migration Assessment"
        core_props.comments = f"Document generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        core_props.created = datetime.now()
        core_props.modified = datetime.now()
    
    def _convert_markdown_content(self, markdown_content: str):
        """Convert markdown content to Word document"""
        
        lines = markdown_content.split('\n')
        in_code_block = False
        in_table = False
        table_data = []
        table_headers = []
        
        for i, line in enumerate(lines):
            line = line.rstrip()
            
            # Handle code blocks
            if line.strip().startswith('```'):
                if in_code_block:
                    # End code block
                    in_code_block = False
                    continue
                else:
                    # Start code block
                    in_code_block = True
                    continue
            
            if in_code_block:
                # Add code line
                p = self.doc.add_paragraph(line, style=self.styles['code'])
                continue
            
            # Handle tables
            if line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                    table_headers = []
                
                # Parse table row
                cells = [cell.strip() for cell in line.split('|')[1:-1]]
                
                # Check if this is a separator line (contains only dashes)
                if all(cell.replace('-', '').replace(':', '').strip() == '' for cell in cells):
                    continue
                
                # Check if this is the header row
                if not table_headers:
                    table_headers = cells
                else:
                    table_data.append(cells)
                
                continue
            else:
                # End table if we were in one
                if in_table and table_data:
                    self._add_table(table_headers, table_data)
                    in_table = False
                    table_data = []
                    table_headers = []
            
            # Handle headings
            if line.startswith('# '):
                self.doc.add_heading(line[2:], level=1).style = self.styles['h1']
            elif line.startswith('## '):
                self.doc.add_heading(line[3:], level=2).style = self.styles['h2']
            elif line.startswith('### '):
                self.doc.add_heading(line[4:], level=3).style = self.styles['h3']
            elif line.startswith('#### '):
                self.doc.add_heading(line[5:], level=4).style = self.styles['h3']
            elif line.startswith('- '):
                # Bullet list
                p = self.doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\. ', line):
                # Numbered list
                p = self.doc.add_paragraph(line[3:], style='List Number')
            elif line.strip():
                # Regular paragraph with formatting
                self._add_formatted_paragraph(line)
            else:
                # Empty line - add spacing
                self.doc.add_paragraph()
        
        # Handle any remaining table
        if in_table and table_data:
            self._add_table(table_headers, table_data)
    
    def _add_formatted_paragraph(self, text: str):
        """Add paragraph with inline formatting"""
        p = self.doc.add_paragraph()
        
        # Split by formatting markers
        parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                # Inline code
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            elif part.startswith('*') and part.endswith('*') and len(part) > 2:
                # Italic text
                run = p.add_run(part[1:-1])
                run.font.italic = True
            else:
                # Regular text
                p.add_run(part)
    
    def _add_table(self, headers: List[str], data: List[List[str]]):
        """Add formatted table to document"""
        if not data:
            return
        
        # Create table
        table = self.doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Add headers
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
            # Add light blue background
            cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="4472C4"/>'))
        
        # Add data rows
        for row_idx, row_data in enumerate(data):
            table_row = table.rows[row_idx + 1]
            for col_idx, cell_data in enumerate(row_data):
                cell = table_row.cells[col_idx]
                cell.text = cell_data
        
        # Add spacing after table
        self.doc.add_paragraph()


def convert_markdown_to_word(markdown_content: str, output_path: str, title: str = "Document") -> str:
    """Convenience function to convert markdown to Word"""
    
    converter = MarkdownToWordConverter()
    return converter.convert_markdown_to_word(markdown_content, output_path, title)


if __name__ == "__main__":
    print("Markdown to Word Converter")
    print("This module converts markdown content to professional Word documents.") 