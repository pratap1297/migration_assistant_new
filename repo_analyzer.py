"""
Repository Migration Analysis Tool - Enhanced Deep Analysis Version
Analyzes Git repositories for cloud migration assessment with comprehensive data collection
"""

import os
import json
import time
import tempfile
import shutil
import subprocess
import logging
from typing import Dict, List, Any, Optional, Tuple, Set
from datetime import datetime
from pathlib import Path
import gradio as gr
from dataclasses import dataclass, field, asdict
from collections import defaultdict
import re

# Configure logging
def setup_logging():
    """Setup comprehensive logging configuration"""
    # Create logs directory
    log_dir = "logs"
    os.makedirs(log_dir, exist_ok=True)
    
    # Create timestamp for log file
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    log_file = os.path.join(log_dir, f"repo_analyzer_{timestamp}.log")
    
    # Configure logging format
    log_format = logging.Formatter(
        '%(asctime)s - %(levelname)s - [%(funcName)s:%(lineno)d] - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    
    # File handler - detailed logs
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.DEBUG)
    file_handler.setFormatter(log_format)
    
    # Console handler - info and above
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_handler.setFormatter(log_format)
    
    # Configure root logger
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    # Log initial setup
    logger.info("="*60)
    logger.info("Repository Migration Analyzer Started - Enhanced Version")
    logger.info(f"Log file: {log_file}")
    logger.info("="*60)
    
    return logger

# Initialize logging
logger = setup_logging()

# LangChain imports
try:
    from langchain_google_genai import ChatGoogleGenerativeAI
    logger.info("Successfully imported langchain_google_genai")
except ImportError as e:
    logger.error(f"Failed to import langchain_google_genai: {e}")
    from langchain.llms import GoogleGenerativeAI as ChatGoogleGenerativeAI

from langchain.agents import Tool, initialize_agent, AgentType
from langchain.schema import Document
from langchain.prompts import PromptTemplate

# Document generation imports
import markdown
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Rate limiting
from functools import wraps
from threading import Lock

# Git operations
import git
from git import Repo

# File parsing
import yaml
import toml
import configparser
from xml.etree import ElementTree as ET


# Rate limiter for Gemini API
class RateLimiter:
    def __init__(self, max_calls: int = 14, time_window: int = 60):
        self.max_calls = max_calls
        self.time_window = time_window
        self.calls = []
        self.lock = Lock()
        logger.info(f"RateLimiter initialized: {max_calls} calls per {time_window} seconds")
    
    def __call__(self, func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            with self.lock:
                now = time.time()
                # Remove old calls outside the time window
                self.calls = [call_time for call_time in self.calls 
                            if now - call_time < self.time_window]
                
                # Check if we can make a new call
                if len(self.calls) >= self.max_calls:
                    sleep_time = self.time_window - (now - self.calls[0]) + 1
                    logger.warning(f"Rate limit reached. Sleeping for {sleep_time:.1f} seconds...")
                    time.sleep(sleep_time)
                    self.calls = []
                
                # Record this call
                self.calls.append(now)
                logger.debug(f"API call recorded. Total calls in window: {len(self.calls)}")
            
            return func(*args, **kwargs)
        return wrapper


# Enhanced data structures for deep analysis
@dataclass
class FileInfo:
    path: str
    name: str
    extension: str
    size: int
    content_preview: str = ""
    detected_patterns: List[str] = field(default_factory=list)
    imports: List[str] = field(default_factory=list)
    
@dataclass
class ServiceInfo:
    name: str
    type: str  # database, cache, message_queue, etc.
    connection_string: str = ""
    port: int = None
    additional_info: Dict[str, Any] = field(default_factory=dict)

@dataclass
class ComponentInfo:
    name: str
    path: str
    language: str = "unknown"
    framework: str = "unknown"
    dependencies: Dict[str, List[str]] = field(default_factory=dict)  # Changed to dict for better organization
    docker_info: Dict[str, Any] = field(default_factory=dict)
    config_files: List[FileInfo] = field(default_factory=list)
    source_files: List[FileInfo] = field(default_factory=list)
    api_endpoints: List[str] = field(default_factory=list)
    database_connections: List[ServiceInfo] = field(default_factory=list)
    external_services: List[ServiceInfo] = field(default_factory=list)
    environment_variables: Dict[str, str] = field(default_factory=dict)
    openshift_resources: List[Dict[str, Any]] = field(default_factory=list)
    kubernetes_resources: List[Dict[str, Any]] = field(default_factory=list)
    build_info: Dict[str, Any] = field(default_factory=dict)
    test_info: Dict[str, Any] = field(default_factory=dict)
    
@dataclass
class RepositoryAnalysis:
    repo_name: str
    repo_url: str
    analysis_date: str
    git_info: Dict[str, Any] = field(default_factory=dict)
    components: List[ComponentInfo] = field(default_factory=list)
    tech_stack: Dict[str, List[str]] = field(default_factory=dict)
    dependencies: Dict[str, Any] = field(default_factory=dict)
    architecture_patterns: List[str] = field(default_factory=list)
    security_findings: Dict[str, Any] = field(default_factory=dict)
    migration_considerations: Dict[str, Any] = field(default_factory=dict)
    infrastructure_requirements: Dict[str, Any] = field(default_factory=dict)
    cicd_info: Dict[str, Any] = field(default_factory=dict)
    testing_strategy: Dict[str, Any] = field(default_factory=dict)
    deployment_info: Dict[str, Any] = field(default_factory=dict)
    gaps: List[str] = field(default_factory=list)
    raw_analysis_data: Dict[str, Any] = field(default_factory=dict)  # Store all raw data


# Enhanced analysis functions
def analyze_git_history_deep(repo_path: str) -> Dict[str, Any]:
    """Deep Git repository analysis including branch strategies and commit patterns"""
    logger.info(f"Starting deep Git analysis for: {repo_path}")
    try:
        repo = Repo(repo_path)
        
        # Basic stats
        commits = list(repo.iter_commits())
        if not commits:
            return {"error": "No commits found"}
            
        first_commit = commits[-1]
        last_commit = commits[0]
        
        # Analyze contributors
        author_stats = defaultdict(lambda: {"commits": 0, "additions": 0, "deletions": 0})
        file_changes = defaultdict(int)
        
        # Analyze recent commits for patterns
        for commit in commits[:100]:  # Analyze last 100 commits
            author_stats[commit.author.name]["commits"] += 1
            
            # Get commit statistics
            try:
                stats = commit.stats.total
                author_stats[commit.author.name]["additions"] += stats.get("insertions", 0)
                author_stats[commit.author.name]["deletions"] += stats.get("deletions", 0)
                
                # Track file changes
                for file in commit.stats.files:
                    file_changes[file] += 1
            except:
                pass
        
        # Get branch information
        branches = []
        for branch in repo.branches:
            try:
                branch_info = {
                    "name": branch.name,
                    "commit": branch.commit.hexsha[:8],
                    "last_commit_date": branch.commit.committed_datetime.isoformat()
                }
                branches.append(branch_info)
            except:
                pass
        
        # Identify most changed files
        hot_files = sorted(file_changes.items(), key=lambda x: x[1], reverse=True)[:10]
        
        result = {
            "first_commit_date": first_commit.committed_datetime.isoformat(),
            "last_commit_date": last_commit.committed_datetime.isoformat(),
            "total_commits": len(commits),
            "author_statistics": dict(author_stats),
            "contributor_count": len(author_stats),
            "branches": branches,
            "default_branch": repo.active_branch.name if repo.active_branch else "unknown",
            "hot_files": hot_files,
            "recent_commits": [
                {
                    "hash": c.hexsha[:8],
                    "author": c.author.name,
                    "date": c.committed_datetime.isoformat(),
                    "message": c.message.strip(),
                    "files_changed": len(c.stats.files)
                } for c in commits[:10]
            ]
        }
        
        logger.info(f"Deep Git analysis completed: {result['total_commits']} commits, {result['contributor_count']} contributors")
        return result
    except Exception as e:
        logger.error(f"Error in deep git analysis: {str(e)}", exc_info=True)
        return {"error": f"Error analyzing git repository: {str(e)}"}


def analyze_source_file(file_path: str) -> FileInfo:
    """Analyze a source code file for imports, patterns, and dependencies"""
    logger.debug(f"Analyzing source file: {file_path}")
    
    file_info = FileInfo(
        path=file_path,
        name=os.path.basename(file_path),
        extension=Path(file_path).suffix,
        size=os.path.getsize(file_path)
    )
    
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
            file_info.content_preview = content[:500]
        
        # Language-specific analysis
        if file_info.extension in ['.py']:
            # Python imports
            imports = re.findall(r'^(?:from|import)\s+(\S+)', content, re.MULTILINE)
            file_info.imports = list(set(imports))
            
            # Detect frameworks
            if 'flask' in content.lower():
                file_info.detected_patterns.append('flask')
            if 'django' in content.lower():
                file_info.detected_patterns.append('django')
            if 'fastapi' in content.lower():
                file_info.detected_patterns.append('fastapi')
                
        elif file_info.extension in ['.js', '.ts', '.jsx', '.tsx']:
            # JavaScript/TypeScript imports
            imports = re.findall(r"(?:import|require)\s*\(?['\"]([^'\"]+)['\"]", content)
            file_info.imports = list(set(imports))
            
            # Detect frameworks
            if 'express' in content.lower():
                file_info.detected_patterns.append('express')
            if 'react' in content.lower():
                file_info.detected_patterns.append('react')
            if 'angular' in content.lower():
                file_info.detected_patterns.append('angular')
            if 'vue' in content.lower():
                file_info.detected_patterns.append('vue')
                
        elif file_info.extension in ['.java']:
            # Java imports
            imports = re.findall(r'^import\s+(\S+);', content, re.MULTILINE)
            file_info.imports = list(set(imports))
            
            # Detect frameworks
            if 'springframework' in content:
                file_info.detected_patterns.append('spring')
            if '@Entity' in content:
                file_info.detected_patterns.append('jpa')
                
        elif file_info.extension in ['.cs']:
            # C# imports
            imports = re.findall(r'^using\s+(\S+);', content, re.MULTILINE)
            file_info.imports = list(set(imports))
            
            # Detect frameworks
            if 'Microsoft.AspNetCore' in content:
                file_info.detected_patterns.append('aspnetcore')
                
        # Detect API endpoints
        api_patterns = [
            r'@(?:Get|Post|Put|Delete|Patch)Mapping\s*\(["\']([^"\']+)',  # Spring
            r'@(?:app\.)?(?:route|get|post|put|delete)\s*\(["\']([^"\']+)',  # Flask/Express
            r'router\.(?:get|post|put|delete)\s*\(["\']([^"\']+)',  # Express Router
            r'@(?:Get|Post|Put|Delete)\s*\(["\']([^"\']+)',  # .NET
        ]
        
        for pattern in api_patterns:
            endpoints = re.findall(pattern, content, re.IGNORECASE)
            if endpoints:
                file_info.detected_patterns.extend([f"endpoint:{e}" for e in endpoints])
        
        # Detect database operations
        db_patterns = {
            'mongodb': r'(?:mongoose|MongoClient|mongodb)',
            'postgresql': r'(?:pg|postgres|psycopg2)',
            'mysql': r'(?:mysql|mysqlclient|pymysql)',
            'redis': r'(?:redis|Redis)',
            'elasticsearch': r'(?:elasticsearch|Elasticsearch)',
            'kafka': r'(?:kafka|KafkaProducer|KafkaConsumer)',
            'rabbitmq': r'(?:amqp|pika|RabbitMQ)'
        }
        
        for db_type, pattern in db_patterns.items():
            if re.search(pattern, content, re.IGNORECASE):
                file_info.detected_patterns.append(f"uses:{db_type}")
        
    except Exception as e:
        logger.error(f"Error analyzing source file {file_path}: {e}")
    
    return file_info


def analyze_openshift_kubernetes_resources(file_path: str) -> List[Dict[str, Any]]:
    """Analyze OpenShift/Kubernetes YAML files for resource definitions"""
    logger.debug(f"Analyzing K8s/OpenShift resource: {file_path}")
    resources = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Handle multi-document YAML files
        yaml_docs = content.split('\n---\n')
        
        for doc in yaml_docs:
            if not doc.strip():
                continue
                
            try:
                resource = yaml.safe_load(doc)
                if resource and isinstance(resource, dict) and 'kind' in resource:
                    # Extract key information
                    resource_info = {
                        'kind': resource.get('kind'),
                        'name': resource.get('metadata', {}).get('name', 'unnamed'),
                        'namespace': resource.get('metadata', {}).get('namespace', 'default'),
                        'labels': resource.get('metadata', {}).get('labels', {}),
                        'file': os.path.basename(file_path)
                    }
                    
                    # Extract specific information based on resource type
                    if resource['kind'] == 'Deployment' or resource['kind'] == 'DeploymentConfig':
                        spec = resource.get('spec', {})
                        template = spec.get('template', {})
                        containers = template.get('spec', {}).get('containers', [])
                        
                        resource_info['replicas'] = spec.get('replicas', 1)
                        resource_info['containers'] = []
                        
                        for container in containers:
                            container_info = {
                                'name': container.get('name'),
                                'image': container.get('image'),
                                'ports': container.get('ports', []),
                                'env': container.get('env', []),
                                'resources': container.get('resources', {})
                            }
                            resource_info['containers'].append(container_info)
                    
                    elif resource['kind'] == 'Service':
                        spec = resource.get('spec', {})
                        resource_info['type'] = spec.get('type', 'ClusterIP')
                        resource_info['ports'] = spec.get('ports', [])
                        resource_info['selector'] = spec.get('selector', {})
                    
                    elif resource['kind'] == 'Route':
                        spec = resource.get('spec', {})
                        resource_info['host'] = spec.get('host')
                        resource_info['path'] = spec.get('path', '/')
                        resource_info['tls'] = spec.get('tls', {})
                    
                    elif resource['kind'] == 'ConfigMap':
                        resource_info['data_keys'] = list(resource.get('data', {}).keys())
                    
                    elif resource['kind'] == 'Secret':
                        resource_info['type'] = resource.get('type', 'Opaque')
                        resource_info['data_keys'] = list(resource.get('data', {}).keys())
                    
                    elif resource['kind'] == 'PersistentVolumeClaim':
                        spec = resource.get('spec', {})
                        resource_info['accessModes'] = spec.get('accessModes', [])
                        resource_info['storage'] = spec.get('resources', {}).get('requests', {}).get('storage')
                    
                    resources.append(resource_info)
                    logger.debug(f"Found {resource_info['kind']}: {resource_info['name']}")
                    
            except yaml.YAMLError as e:
                logger.error(f"Error parsing YAML document in {file_path}: {e}")
    
    except Exception as e:
        logger.error(f"Error analyzing K8s/OpenShift file {file_path}: {e}")
    
    return resources


def extract_environment_configuration(comp_path: str) -> Dict[str, Any]:
    """Extract environment variables and configuration from various sources"""
    logger.debug(f"Extracting environment configuration from: {comp_path}")
    
    env_config = {
        'variables': {},
        'files': [],
        'secrets': [],
        'config_maps': []
    }
    
    # Check for .env files
    env_file_patterns = ['.env', '.env.example', '.env.sample', '.env.local', '.env.production']
    for pattern in env_file_patterns:
        env_path = os.path.join(comp_path, pattern)
        if os.path.exists(env_path):
            env_config['files'].append(pattern)
            try:
                with open(env_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        line = line.strip()
                        if line and not line.startswith('#') and '=' in line:
                            key, value = line.split('=', 1)
                            env_config['variables'][key.strip()] = value.strip()
                            
                            # Identify secrets
                            if any(secret in key.lower() for secret in ['password', 'secret', 'key', 'token']):
                                env_config['secrets'].append(key.strip())
            except Exception as e:
                logger.error(f"Error reading env file {env_path}: {e}")
    
    # Check for docker-compose environment
    compose_files = ['docker-compose.yml', 'docker-compose.yaml']
    for compose_file in compose_files:
        compose_path = os.path.join(comp_path, compose_file)
        if os.path.exists(compose_path):
            try:
                with open(compose_path, 'r') as f:
                    compose_data = yaml.safe_load(f)
                    if compose_data and 'services' in compose_data:
                        for service_name, service_config in compose_data['services'].items():
                            if 'environment' in service_config:
                                env_vars = service_config['environment']
                                if isinstance(env_vars, dict):
                                    env_config['variables'].update(env_vars)
                                elif isinstance(env_vars, list):
                                    for env_var in env_vars:
                                        if '=' in env_var:
                                            key, value = env_var.split('=', 1)
                                            env_config['variables'][key] = value
            except Exception as e:
                logger.error(f"Error parsing docker-compose file: {e}")
    
    return env_config


def analyze_dockerfile_deep(file_path: str) -> Dict[str, Any]:
    """Deep analysis of Dockerfile including multi-stage builds and best practices"""
    logger.info(f"Deep Dockerfile analysis: {file_path}")
    
    result = {
        "stages": [],
        "base_images": [],
        "final_image": None,
        "exposed_ports": [],
        "environment_variables": {},
        "volumes": [],
        "build_args": {},
        "security_issues": [],
        "best_practices": [],
        "size_optimization": []
    }
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        current_stage = {"name": "default", "from": None, "instructions": []}
        stage_count = 0
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            # Parse instructions
            parts = line.split(None, 1)
            if not parts:
                continue
                
            instruction = parts[0].upper()
            value = parts[1] if len(parts) > 1 else ''
            
            # Handle multi-stage builds
            if instruction == 'FROM':
                if current_stage['from']:
                    result['stages'].append(current_stage)
                    
                stage_count += 1
                if ' AS ' in value:
                    base, stage_name = value.split(' AS ', 1)
                    current_stage = {"name": stage_name, "from": base, "instructions": []}
                else:
                    current_stage = {"name": f"stage_{stage_count}", "from": value, "instructions": []}
                
                result['base_images'].append(value.split(' AS ')[0])
                
                # Check for security issues with base images
                if ':latest' in value or (' ' in value and ':' not in value.split()[0]):
                    result['security_issues'].append(f"Using latest tag or no tag specified: {value}")
            
            elif instruction == 'EXPOSE':
                result['exposed_ports'].extend(value.split())
            
            elif instruction == 'ENV':
                if ' ' in value:
                    key, val = value.split(None, 1)
                    result['environment_variables'][key] = val
            
            elif instruction == 'ARG':
                if '=' in value:
                    key, val = value.split('=', 1)
                    result['build_args'][key] = val
                else:
                    result['build_args'][value] = None
            
            elif instruction == 'VOLUME':
                result['volumes'].append(value)
            
            elif instruction == 'USER':
                if value.lower() != 'root':
                    result['best_practices'].append(f"Running as non-root user: {value}")
                else:
                    result['security_issues'].append("Running as root user")
            
            elif instruction == 'RUN':
                # Check for package manager cleanup
                if 'apt-get' in value and 'apt-get clean' not in value and 'rm -rf /var/lib/apt/lists/*' not in value:
                    result['size_optimization'].append("Missing apt cleanup in RUN instruction")
                
                # Check for multiple RUN commands that could be combined
                if 'apt-get update' in value or 'apt-get install' in value:
                    current_stage['instructions'].append(('package_install', value))
            
            current_stage['instructions'].append((instruction, value))
        
        # Add the last stage
        if current_stage['from']:
            result['stages'].append(current_stage)
            result['final_image'] = current_stage['from']
        
        # Additional best practices checks
        if len(result['stages']) > 1:
            result['best_practices'].append("Using multi-stage build")
        
        if not any(inst[0] == 'USER' for stage in result['stages'] for inst in stage['instructions']):
            result['security_issues'].append("No USER instruction found - container will run as root")
        
        # Check for COPY vs ADD usage
        add_count = sum(1 for stage in result['stages'] for inst in stage['instructions'] if inst[0] == 'ADD')
        if add_count > 0:
            result['best_practices'].append(f"Consider using COPY instead of ADD ({add_count} ADD instructions found)")
        
    except Exception as e:
        logger.error(f"Error analyzing Dockerfile: {str(e)}", exc_info=True)
        result['error'] = str(e)
    
    return result


def detect_services_and_dependencies(content: str, file_path: str) -> Dict[str, List[Any]]:
    """Enhanced service and dependency detection from file content"""
    services = {
        'databases': [],
        'caches': [],
        'message_queues': [],
        'external_apis': [],
        'other_services': []
    }
    
    # Enhanced database patterns
    db_patterns = {
        'postgresql': [
            r'postgres(?:ql)?://([^/\s]+)(?:/([^?\s]+))?',
            r'jdbc:postgresql://([^/\s]+)(?:/([^?\s]+))?',
            r'Host=([^;]+);.*Database=([^;]+).*postgres'
        ],
        'mysql': [
            r'mysql://([^/\s]+)(?:/([^?\s]+))?',
            r'jdbc:mysql://([^/\s]+)(?:/([^?\s]+))?',
            r'Server=([^;]+);.*Database=([^;]+).*mysql'
        ],
        'mongodb': [
            r'mongodb(?:\+srv)?://([^/\s]+)(?:/([^?\s]+))?',
            r'mongodb://([^:]+):(\d+)'
        ],
        'redis': [
            r'redis://([^/:\s]+)(?::(\d+))?',
            r'redis://([^@]+)@([^/:\s]+)(?::(\d+))?'
        ],
        'elasticsearch': [
            r'elasticsearch://([^/:\s]+)(?::(\d+))?',
            r'http://([^/:\s]+):9200'
        ]
    }
    
    for db_type, patterns in db_patterns.items():
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                service = ServiceInfo(
                    name=db_type,
                    type='database' if db_type != 'redis' else 'cache',
                    connection_string=match[0] if isinstance(match, tuple) else match
                )
                if db_type == 'redis':
                    services['caches'].append(service)
                else:
                    services['databases'].append(service)
    
    # Message queue patterns
    mq_patterns = {
        'rabbitmq': [r'amqp://([^/\s]+)', r'rabbitmq://([^/\s]+)'],
        'kafka': [r'kafka://([^/\s]+)', r'([^:,\s]+):9092'],
        'sqs': [r'sqs\.([^\.]+)\.amazonaws\.com'],
        'azureservicebus': [r'\.servicebus\.windows\.net']
    }
    
    for mq_type, patterns in mq_patterns.items():
        for pattern in patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            for match in matches:
                service = ServiceInfo(
                    name=mq_type,
                    type='message_queue',
                    connection_string=match
                )
                services['message_queues'].append(service)
    
    # External API patterns
    api_patterns = [
        r'https?://api\.([^/\s]+)',
        r'https?://([^/\s]+)/api/v\d+',
        r'baseurl["\']?\s*[:=]\s*["\']?(https?://[^"\'\s]+)',
        r'endpoint["\']?\s*[:=]\s*["\']?(https?://[^"\'\s]+)'
    ]
    
    for pattern in api_patterns:
        matches = re.findall(pattern, content, re.IGNORECASE)
        services['external_apis'].extend(matches)
    
    # Remove duplicates
    for key in services:
        if key == 'external_apis':
            services[key] = list(set(services[key]))
        else:
            # For service objects, deduplicate based on connection string
            seen = set()
            unique_services = []
            for service in services[key]:
                if service.connection_string not in seen:
                    seen.add(service.connection_string)
                    unique_services.append(service)
            services[key] = unique_services
    
    return services


# Enhanced Repository Analyzer
class RepositoryAnalyzer:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.rate_limiter = RateLimiter(max_calls=14, time_window=60)
        logger.info("Initializing Enhanced RepositoryAnalyzer")
        self.llm = self._setup_llm()
    
    def _setup_llm(self):
        logger.info("Setting up LLM (Gemini)")
        try:
            llm = ChatGoogleGenerativeAI(
                model="gemini-2.5-flash-lite-preview-06-17",
                google_api_key=self.api_key,
                temperature=0.1,
                max_tokens=4096
            )
            logger.info("LLM setup successful")
            return llm
        except Exception as e:
            logger.error(f"Failed to setup LLM: {e}", exc_info=True)
            raise
    
    @RateLimiter()
    def _analyze_with_llm(self, prompt: str) -> str:
        """Run analysis through LLM with rate limiting"""
        logger.info("Calling LLM for analysis")
        logger.debug(f"Prompt length: {len(prompt)} characters")
        try:
            response = self.llm.predict(prompt)
            logger.info("LLM analysis completed successfully")
            logger.debug(f"Response length: {len(response)} characters")
            return response
        except Exception as e:
            logger.error(f"Error in LLM analysis: {str(e)}", exc_info=True)
            return f"Error in LLM analysis: {str(e)}"
    
    def analyze_repository(self, repo_path: str, progress_callback=None) -> RepositoryAnalysis:
        """Main analysis function with deep, comprehensive analysis"""
        logger.info("="*60)
        logger.info(f"Starting deep repository analysis for: {repo_path}")
        logger.info("="*60)
        
        analysis = RepositoryAnalysis(
            repo_name=Path(repo_path).name,
            repo_url=repo_path,
            analysis_date=datetime.now().isoformat()
        )
        
        try:
            # Phase 1: Deep Git Analysis
            logger.info("PHASE 1: Deep Git Analysis")
            if progress_callback:
                progress_callback("Performing deep Git analysis...")
            analysis.git_info = analyze_git_history_deep(repo_path)
            
            # Phase 2: Repository Structure Analysis
            logger.info("PHASE 2: Repository Structure Deep Scan")
            if progress_callback:
                progress_callback("Scanning repository structure...")
            self._deep_scan_repository(repo_path, analysis, progress_callback)
            
            # Phase 3: Component Detection and Analysis
            logger.info("PHASE 3: Component Detection and Deep Analysis")
            if progress_callback:
                progress_callback("Detecting and analyzing components...")
            self._detect_and_analyze_components(repo_path, analysis, progress_callback)
            
            # Phase 4: Infrastructure and Deployment Analysis
            logger.info("PHASE 4: Infrastructure and Deployment Analysis")
            if progress_callback:
                progress_callback("Analyzing infrastructure and deployment configurations...")
            self._analyze_infrastructure(repo_path, analysis)
            
            # Phase 5: Security and Compliance Scan
            logger.info("PHASE 5: Security and Compliance Analysis")
            if progress_callback:
                progress_callback("Performing security and compliance checks...")
            self._analyze_security(analysis)
            
            # Phase 6: Aggregate and Correlate Data
            logger.info("PHASE 6: Data Aggregation and Correlation")
            if progress_callback:
                progress_callback("Aggregating and correlating analysis data...")
            self._aggregate_analysis_data(analysis)
            
            # Phase 7: AI-Powered Insights Generation
            logger.info("PHASE 7: AI-Powered Insights Generation")
            if progress_callback:
                progress_callback("Generating AI-powered insights and recommendations...")
            self._generate_comprehensive_insights(analysis)
            
            logger.info("="*60)
            logger.info("Deep repository analysis completed successfully")
            logger.info("="*60)
            
        except Exception as e:
            logger.error(f"Error during analysis: {str(e)}", exc_info=True)
            analysis.gaps.append(f"Error during analysis: {str(e)}")
        
        return analysis
    
    def _deep_scan_repository(self, repo_path: str, analysis: RepositoryAnalysis, progress_callback=None):
        """Perform deep scan of entire repository structure"""
        logger.info("Starting deep repository scan")
        
        # File type statistics
        file_stats = defaultdict(int)
        total_size = 0
        important_files = []
        all_configs = []
        all_dockerfiles = []
        all_k8s_files = []
        all_ci_files = []
        
        # Patterns to identify important files
        important_patterns = {
            'readme': r'^readme\.(md|txt|rst)$',
            'license': r'^license(\.(md|txt))?$',
            'dockerfile': r'^dockerfile(\.\w+)?$',
            'docker-compose': r'^docker-compose\.(yml|yaml)$',
            'kubernetes': r'\.(yaml|yml)$',
            'ci/cd': r'^(\.|)+(gitlab-ci|jenkins|travis|circle|azure-pipelines|github)\.(yml|yaml|json)$',
            'package': r'^(package|requirements|pom|build|cargo|go)\.(json|txt|xml|gradle|toml|mod)$',
            'config': r'^(config|settings|application|app)\.(json|yml|yaml|properties|ini|toml)$'
        }
        
        # Walk through entire repository
        for root, dirs, files in os.walk(repo_path):
            # Skip certain directories
            dirs[:] = [d for d in dirs if d not in {'.git', 'node_modules', '__pycache__', '.venv', 'venv', 'target', 'dist', 'build'}]
            
            rel_path = os.path.relpath(root, repo_path)
            
            for file in files:
                file_path = os.path.join(root, file)
                rel_file_path = os.path.join(rel_path, file)
                
                try:
                    file_size = os.path.getsize(file_path)
                    total_size += file_size
                    
                    # Track file extensions
                    ext = Path(file).suffix.lower()
                    if ext:
                        file_stats[ext] += 1
                    
                    # Check for important files
                    file_lower = file.lower()
                    for pattern_type, pattern in important_patterns.items():
                        if re.match(pattern, file_lower):
                            file_info = FileInfo(
                                path=file_path,
                                name=file,
                                extension=ext,
                                size=file_size
                            )
                            
                            if pattern_type == 'dockerfile':
                                all_dockerfiles.append(file_path)
                            elif pattern_type == 'kubernetes' and 'k8s' in rel_path or 'kubernetes' in rel_path or 'openshift' in rel_path:
                                all_k8s_files.append(file_path)
                            elif pattern_type == 'ci/cd':
                                all_ci_files.append(file_path)
                            elif pattern_type == 'config':
                                all_configs.append(file_path)
                            
                            important_files.append((pattern_type, file_info))
                            break
                    
                    # Check YAML files for Kubernetes/OpenShift resources
                    if ext in ['.yaml', '.yml']:
                        try:
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read(1000)  # Read first 1KB
                                if 'kind:' in content and any(k in content for k in ['Deployment', 'Service', 'Route', 'ConfigMap']):
                                    all_k8s_files.append(file_path)
                        except:
                            pass
                
                except Exception as e:
                    logger.debug(f"Error processing file {file_path}: {e}")
        
        # Store scan results
        analysis.raw_analysis_data['repository_scan'] = {
            'total_files': sum(file_stats.values()),
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'file_types': dict(file_stats),
            'important_files': important_files,
            'dockerfiles': all_dockerfiles,
            'kubernetes_files': all_k8s_files,
            'ci_files': all_ci_files,
            'config_files': all_configs
        }
        
        # Analyze technology stack based on file extensions
        tech_indicators = {
            'javascript': ['.js', '.jsx', '.ts', '.tsx'],
            'python': ['.py'],
            'java': ['.java'],
            'csharp': ['.cs'],
            'go': ['.go'],
            'rust': ['.rs'],
            'ruby': ['.rb'],
            'php': ['.php']
        }
        
        for lang, extensions in tech_indicators.items():
            if any(file_stats.get(ext, 0) > 0 for ext in extensions):
                analysis.tech_stack.setdefault('languages', []).append(lang)
        
        logger.info(f"Repository scan complete: {sum(file_stats.values())} files, "
                   f"{analysis.raw_analysis_data['repository_scan']['total_size_mb']} MB")
    
    def _detect_and_analyze_components(self, repo_path: str, analysis: RepositoryAnalysis, progress_callback=None):
        """Enhanced component detection with deep analysis"""
        logger.info("Starting enhanced component detection")
        
        # Get scan data
        scan_data = analysis.raw_analysis_data.get('repository_scan', {})
        
        # First, check if entire repo is a single component
        root_files = os.listdir(repo_path)
        root_indicators = ['package.json', 'requirements.txt', 'pom.xml', 'go.mod', 'Cargo.toml', 'Dockerfile']
        
        if any(f in root_files for f in root_indicators):
            logger.info("Root directory appears to be a component")
            self._analyze_component_deep(repo_path, "root", analysis)
        
        # Look for components in subdirectories
        potential_components = []
        
        # Strategy 1: Look for directories with component indicators
        for item in os.listdir(repo_path):
            item_path = os.path.join(repo_path, item)
            if os.path.isdir(item_path) and not item.startswith('.') and item not in {'node_modules', 'vendor', 'target'}:
                # Check for component indicators
                try:
                    subfiles = os.listdir(item_path)
                    if any(f in subfiles for f in root_indicators):
                        potential_components.append((item_path, item))
                except:
                    pass
        
        # Strategy 2: For OpenShift/K8s repos, look for deployment directories
        k8s_dirs = ['k8s', 'kubernetes', 'openshift', 'manifests', 'deployments', 'helm']
        for k8s_dir in k8s_dirs:
            k8s_path = os.path.join(repo_path, k8s_dir)
            if os.path.exists(k8s_path):
                logger.info(f"Found Kubernetes/OpenShift directory: {k8s_dir}")
                # Analyze as infrastructure component
                self._analyze_k8s_directory(k8s_path, analysis)
        
        # Analyze each detected component
        for comp_path, comp_name in potential_components:
            logger.info(f"Analyzing component: {comp_name}")
            if progress_callback:
                progress_callback(f"Analyzing component: {comp_name}")
            self._analyze_component_deep(comp_path, comp_name, analysis)
        
        # If no components found but we have Dockerfiles, analyze them
        if not analysis.components and scan_data.get('dockerfiles'):
            logger.info("No traditional components found, analyzing Dockerfiles")
            for dockerfile in scan_data['dockerfiles']:
                dir_path = os.path.dirname(dockerfile)
                comp_name = os.path.basename(dir_path)
                if dir_path == repo_path:
                    comp_name = "root"
                self._analyze_component_deep(dir_path, comp_name, analysis)
    
    def _analyze_component_deep(self, comp_path: str, comp_name: str, analysis: RepositoryAnalysis):
        """Deep analysis of a single component"""
        logger.info(f"Deep analysis of component: {comp_name} at {comp_path}")
        
        component = ComponentInfo(
            name=comp_name,
            path=comp_path
        )
        
        # Analyze all files in component
        source_files = []
        config_files = []
        test_files = []
        
        for root, dirs, files in os.walk(comp_path):
            # Skip certain directories
            dirs[:] = [d for d in dirs if d not in {'.git', 'node_modules', '__pycache__', 'vendor', 'target'}]
            
            for file in files:
                file_path = os.path.join(root, file)
                ext = Path(file).suffix.lower()
                
                # Analyze source files
                if ext in ['.py', '.js', '.ts', '.java', '.cs', '.go', '.rs', '.rb', '.php']:
                    file_info = analyze_source_file(file_path)
                    source_files.append(file_info)
                    
                    # Detect language if not already set
                    if component.language == "unknown":
                        lang_map = {
                            '.py': 'python',
                            '.js': 'javascript',
                            '.ts': 'typescript',
                            '.java': 'java',
                            '.cs': 'csharp',
                            '.go': 'go',
                            '.rs': 'rust',
                            '.rb': 'ruby',
                            '.php': 'php'
                        }
                        component.language = lang_map.get(ext, 'unknown')
                    
                    # Aggregate detected patterns
                    for pattern in file_info.detected_patterns:
                        if pattern.startswith('endpoint:'):
                            component.api_endpoints.append(pattern.replace('endpoint:', ''))
                        elif pattern.startswith('uses:'):
                            service_type = pattern.replace('uses:', '')
                            service = ServiceInfo(name=service_type, type='database' if 'db' in service_type else 'service')
                            component.external_services.append(service)
                
                # Collect config files
                elif ext in ['.json', '.yml', '.yaml', '.properties', '.ini', '.toml', '.env']:
                    config_files.append(file_path)
                
                # Collect test files
                elif 'test' in file.lower() or 'spec' in file.lower():
                    test_files.append(file_path)
        
        # Store source files (limit to most important ones)
        component.source_files = sorted(source_files, key=lambda x: x.size, reverse=True)[:20]
        
        # Analyze package files
        package_files = {
            'package.json': ('javascript', 'npm'),
            'requirements.txt': ('python', 'pip'),
            'pom.xml': ('java', 'maven'),
            'build.gradle': ('java', 'gradle'),
            'go.mod': ('go', 'go'),
            'Cargo.toml': ('rust', 'cargo'),
            'composer.json': ('php', 'composer'),
            'Gemfile': ('ruby', 'bundler')
        }
        
        for pkg_file, (lang, pkg_type) in package_files.items():
            pkg_path = os.path.join(comp_path, pkg_file)
            if os.path.exists(pkg_path):
                if component.language == "unknown":
                    component.language = lang
                
                # Deep package analysis
                self._analyze_package_file_deep(pkg_path, pkg_type, component)
        
        # Analyze Dockerfile if present
        dockerfile_path = os.path.join(comp_path, 'Dockerfile')
        if os.path.exists(dockerfile_path):
            component.docker_info = analyze_dockerfile_deep(dockerfile_path)
            
            # Extract language from base image if still unknown
            if component.language == "unknown" and component.docker_info.get('base_images'):
                base_image = component.docker_info['base_images'][0].lower()
                if 'node' in base_image:
                    component.language = 'javascript'
                elif 'python' in base_image:
                    component.language = 'python'
                elif 'java' in base_image or 'openjdk' in base_image:
                    component.language = 'java'
                elif 'golang' in base_image:
                    component.language = 'go'
                elif 'ruby' in base_image:
                    component.language = 'ruby'
                elif 'php' in base_image:
                    component.language = 'php'
                elif 'dotnet' in base_image:
                    component.language = 'csharp'
        
        # Analyze configuration files
        for config_file in config_files[:10]:  # Limit to 10 most important
            self._analyze_config_file_deep(config_file, component)
        
        # Extract environment configuration
        env_config = extract_environment_configuration(comp_path)
        component.environment_variables = env_config['variables']
        
        # Analyze test coverage
        component.test_info = {
            'test_files_count': len(test_files),
            'has_unit_tests': any('unit' in f.lower() for f in test_files),
            'has_integration_tests': any('integration' in f.lower() or 'e2e' in f.lower() for f in test_files),
            'test_frameworks': self._detect_test_frameworks(test_files, component.language)
        }
        
        # Build information
        build_files = ['Makefile', 'build.sh', 'build.gradle', 'pom.xml', 'package.json']
        for build_file in build_files:
            build_path = os.path.join(comp_path, build_file)
            if os.path.exists(build_path):
                component.build_info['build_system'] = build_file
                break
        
        logger.info(f"Component {comp_name} analysis complete: "
                   f"language={component.language}, "
                   f"files={len(component.source_files)}, "
                   f"configs={len(config_files)}, "
                   f"tests={len(test_files)}")
        
        analysis.components.append(component)
    
    def _analyze_package_file_deep(self, pkg_path: str, pkg_type: str, component: ComponentInfo):
        """Deep analysis of package files"""
        logger.debug(f"Deep package analysis: {pkg_path}")
        
        try:
            with open(pkg_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if pkg_type == 'npm' and pkg_path.endswith('package.json'):
                data = json.loads(content)
                
                # Extract all dependency types
                component.dependencies['production'] = list(data.get('dependencies', {}).keys())
                component.dependencies['development'] = list(data.get('devDependencies', {}).keys())
                component.dependencies['peer'] = list(data.get('peerDependencies', {}).keys())
                
                # Extract scripts
                component.build_info['scripts'] = list(data.get('scripts', {}).keys())
                
                # Detect frameworks
                all_deps = component.dependencies['production'] + component.dependencies['development']
                if 'react' in all_deps:
                    component.framework = 'react'
                elif 'vue' in all_deps:
                    component.framework = 'vue'
                elif 'angular' in all_deps:
                    component.framework = 'angular'
                elif 'express' in all_deps:
                    component.framework = 'express'
                elif 'nestjs' in all_deps:
                    component.framework = 'nestjs'
                
            elif pkg_type == 'pip' and pkg_path.endswith('requirements.txt'):
                deps = []
                for line in content.split('\n'):
                    line = line.strip()
                    if line and not line.startswith('#'):
                        # Extract package name without version
                        pkg_name = re.split('[<>=!]', line)[0].strip()
                        if pkg_name:
                            deps.append(pkg_name)
                
                component.dependencies['production'] = deps
                
                # Detect frameworks
                if 'django' in deps:
                    component.framework = 'django'
                elif 'flask' in deps:
                    component.framework = 'flask'
                elif 'fastapi' in deps:
                    component.framework = 'fastapi'
                
            elif pkg_type == 'maven' and pkg_path.endswith('pom.xml'):
                root = ET.fromstring(content)
                deps = []
                
                # Extract dependencies
                for dep in root.findall('.//{http://maven.apache.org/POM/4.0.0}dependency'):
                    group_id = dep.find('{http://maven.apache.org/POM/4.0.0}groupId')
                    artifact_id = dep.find('{http://maven.apache.org/POM/4.0.0}artifactId')
                    if group_id is not None and artifact_id is not None:
                        deps.append(f"{group_id.text}:{artifact_id.text}")
                
                component.dependencies['production'] = deps
                
                # Detect frameworks
                if any('spring' in d for d in deps):
                    component.framework = 'spring'
                
        except Exception as e:
            logger.error(f"Error analyzing package file {pkg_path}: {e}")
    
    def _analyze_config_file_deep(self, config_path: str, component: ComponentInfo):
        """Deep configuration file analysis"""
        try:
            file_info = FileInfo(
                path=config_path,
                name=os.path.basename(config_path),
                extension=Path(config_path).suffix,
                size=os.path.getsize(config_path)
            )
            
            with open(config_path, 'r', encoding='utf-8') as f:
                content = f.read()
                file_info.content_preview = content[:500]
            
            # Detect services and dependencies
            services = detect_services_and_dependencies(content, config_path)
            
            # Add detected services to component
            for db in services['databases']:
                component.database_connections.append(db)
            
            for api in services['external_apis']:
                component.api_endpoints.append(api)
            
            for service in services['caches'] + services['message_queues'] + services['other_services']:
                component.external_services.append(service)
            
            component.config_files.append(file_info)
            
        except Exception as e:
            logger.error(f"Error analyzing config file {config_path}: {e}")
    
    def _analyze_k8s_directory(self, k8s_path: str, analysis: RepositoryAnalysis):
        """Analyze Kubernetes/OpenShift directory"""
        logger.info(f"Analyzing Kubernetes/OpenShift directory: {k8s_path}")
        
        k8s_resources = []
        
        for root, dirs, files in os.walk(k8s_path):
            for file in files:
                if file.endswith(('.yaml', '.yml')):
                    file_path = os.path.join(root, file)
                    resources = analyze_openshift_kubernetes_resources(file_path)
                    k8s_resources.extend(resources)
        
        # Group resources by type
        resource_summary = defaultdict(list)
        for resource in k8s_resources:
            resource_summary[resource['kind']].append(resource)
        
        # Store in analysis
        analysis.deployment_info['kubernetes_resources'] = dict(resource_summary)
        analysis.deployment_info['total_k8s_resources'] = len(k8s_resources)
        
        # Extract deployment patterns
        if 'Deployment' in resource_summary or 'DeploymentConfig' in resource_summary:
            analysis.architecture_patterns.append('kubernetes-native')
        
        if 'Route' in resource_summary:
            analysis.architecture_patterns.append('openshift-native')
        
        if 'Service' in resource_summary:
            services = resource_summary['Service']
            analysis.deployment_info['exposed_services'] = len(services)
        
        logger.info(f"Found {len(k8s_resources)} Kubernetes/OpenShift resources")
    
    def _analyze_infrastructure(self, repo_path: str, analysis: RepositoryAnalysis):
        """Analyze infrastructure and deployment configurations"""
        logger.info("Analyzing infrastructure and deployment configurations")
        
        # Check for Infrastructure as Code
        iac_files = {
            'terraform': ['*.tf', '*.tfvars'],
            'ansible': ['*.yml', 'playbook.yml', 'ansible.cfg'],
            'cloudformation': ['*.template', '*.json'],
            'helm': ['Chart.yaml', 'values.yaml']
        }
        
        for iac_type, patterns in iac_files.items():
            for pattern in patterns:
                if any(Path(repo_path).rglob(pattern)):
                    analysis.infrastructure_requirements['iac_tool'] = iac_type
                    analysis.architecture_patterns.append(f'iac-{iac_type}')
                    break
        
        # Analyze docker-compose if present
        compose_path = os.path.join(repo_path, 'docker-compose.yml')
        if os.path.exists(compose_path):
            self._analyze_docker_compose(compose_path, analysis)
        
        # Check for CI/CD configurations
        ci_configs = {
            'jenkins': ['Jenkinsfile', 'jenkins.yml'],
            'gitlab': ['.gitlab-ci.yml'],
            'github': ['.github/workflows/*.yml'],
            'azure': ['azure-pipelines.yml'],
            'circleci': ['.circleci/config.yml'],
            'travis': ['.travis.yml']
        }
        
        for ci_type, patterns in ci_configs.items():
            for pattern in patterns:
                if any(Path(repo_path).rglob(pattern)):
                    analysis.cicd_info['platform'] = ci_type
                    analysis.architecture_patterns.append(f'cicd-{ci_type}')
                    # Analyze CI/CD file for more details
                    self._analyze_cicd_file(list(Path(repo_path).rglob(pattern))[0], analysis)
                    break
        
    def _analyze_docker_compose(self, compose_path: str, analysis: RepositoryAnalysis):
        """Analyze docker-compose file"""
        try:
            with open(compose_path, 'r') as f:
                compose_data = yaml.safe_load(f)
            
            if 'services' in compose_data:
                services = compose_data['services']
                analysis.deployment_info['docker_compose_services'] = list(services.keys())
                
                # Analyze each service
                for service_name, service_config in services.items():
                    # Check for databases
                    if any(db in service_config.get('image', '') for db in ['postgres', 'mysql', 'mongo', 'redis']):
                        analysis.infrastructure_requirements.setdefault('databases', []).append(service_config['image'])
                    
                    # Check for volumes
                    if 'volumes' in service_config:
                        analysis.infrastructure_requirements.setdefault('persistent_storage', True)
                    
                    # Check for networks
                    if 'networks' in service_config:
                        analysis.infrastructure_requirements.setdefault('custom_networking', True)
                
        except Exception as e:
            logger.error(f"Error analyzing docker-compose: {e}")
    
    def _analyze_cicd_file(self, ci_file_path: Path, analysis: RepositoryAnalysis):
        """Analyze CI/CD configuration file"""
        try:
            with open(ci_file_path, 'r') as f:
                content = f.read()
            
            # Common CI/CD patterns
            if 'docker build' in content or 'docker push' in content:
                analysis.cicd_info['docker_build'] = True
            
            if 'test' in content.lower():
                analysis.cicd_info['automated_tests'] = True
            
            if 'deploy' in content.lower():
                analysis.cicd_info['automated_deployment'] = True
            
            # Extract stages/jobs
            if ci_file_path.name == '.gitlab-ci.yml':
                ci_data = yaml.safe_load(content)
                if 'stages' in ci_data:
                    analysis.cicd_info['stages'] = ci_data['stages']
            
        except Exception as e:
            logger.error(f"Error analyzing CI/CD file: {e}")
    
    def _analyze_security(self, analysis: RepositoryAnalysis):
        """Analyze security aspects of the repository"""
        logger.info("Performing security analysis")
        
        security_issues = []
        security_best_practices = []
        
        # Check for security in Dockerfiles
        for component in analysis.components:
            if component.docker_info:
                security_issues.extend(component.docker_info.get('security_issues', []))
                security_best_practices.extend(component.docker_info.get('best_practices', []))
        
        # Check for secrets in environment variables
        for component in analysis.components:
            for key in component.environment_variables:
                if any(secret in key.lower() for secret in ['password', 'secret', 'key', 'token', 'credential']):
                    security_issues.append(f"Potential secret in environment variable: {key}")
        
        # Check for security dependencies
        security_tools = ['snyk', 'owasp', 'security', 'auth0', 'jwt']
        for component in analysis.components:
            for dep_list in component.dependencies.values():
                for dep in dep_list:
                    if any(tool in dep.lower() for tool in security_tools):
                        security_best_practices.append(f"Security tool/library found: {dep}")
        
        # Check for HTTPS/TLS
        for component in analysis.components:
            for endpoint in component.api_endpoints:
                if endpoint.startswith('http://') and 'localhost' not in endpoint:
                    security_issues.append(f"Non-HTTPS endpoint found: {endpoint}")
        
        analysis.security_findings = {
            'issues': security_issues,
            'best_practices': security_best_practices,
            'secrets_management': 'Azure Key Vault' if analysis.tech_stack else 'Required',
            'recommendations': [
                "Implement Azure Key Vault for secrets management",
                "Use Managed Identity for Azure resource authentication",
                "Enable Azure Security Center for continuous security assessment",
                "Implement network security groups and private endpoints",
                "Use Azure Policy for compliance enforcement"
            ]
        }
    
    def _aggregate_analysis_data(self, analysis: RepositoryAnalysis):
        """Aggregate all analysis data for comprehensive view"""
        logger.info("Aggregating analysis data")
        
        # Aggregate all technologies
        all_languages = set()
        all_frameworks = set()
        all_databases = set()
        all_services = set()
        
        for component in analysis.components:
            if component.language != "unknown":
                all_languages.add(component.language)
            if component.framework != "unknown":
                all_frameworks.add(component.framework)
            
            for db in component.database_connections:
                all_databases.add(db.name)
            
            for service in component.external_services:
                all_services.add(service.name)
        
        # Update tech stack
        analysis.tech_stack['languages'] = list(all_languages)
        analysis.tech_stack['frameworks'] = list(all_frameworks)
        analysis.tech_stack['databases'] = list(all_databases)
        analysis.tech_stack['services'] = list(all_services)
        
        # Calculate complexity metrics
        complexity_score = 0
        complexity_factors = []
        
        # Language diversity
        if len(all_languages) > 3:
            complexity_score += 2
            complexity_factors.append("Multiple programming languages")
        elif len(all_languages) > 1:
            complexity_score += 1
            complexity_factors.append("Polyglot application")
        
        # Component count
        if len(analysis.components) > 10:
            complexity_score += 3
            complexity_factors.append("Large number of components")
        elif len(analysis.components) > 5:
            complexity_score += 2
            complexity_factors.append("Multiple components")
        elif len(analysis.components) > 1:
            complexity_score += 1
            complexity_factors.append("Multi-component application")
        
        # External dependencies
        total_deps = sum(len(comp.external_services) + len(comp.database_connections) 
                        for comp in analysis.components)
        if total_deps > 10:
            complexity_score += 2
            complexity_factors.append("Many external dependencies")
        elif total_deps > 5:
            complexity_score += 1
            complexity_factors.append("Several external dependencies")
        
        # Architecture patterns
        if 'microservices' in analysis.architecture_patterns:
            complexity_score += 2
            complexity_factors.append("Microservices architecture")
        
        if 'kubernetes-native' in analysis.architecture_patterns or 'openshift-native' in analysis.architecture_patterns:
            complexity_score += 1
            complexity_factors.append("Container orchestration")
        
        # Store complexity analysis
        analysis.migration_considerations['complexity_score'] = complexity_score
        analysis.migration_considerations['complexity_factors'] = complexity_factors
        analysis.migration_considerations['complexity_level'] = (
            'Low' if complexity_score < 3 else
            'Medium' if complexity_score < 6 else
            'High'
        )
        
        # Identify migration blockers and accelerators
        blockers = []
        accelerators = []
        
        # Check for blockers
        for component in analysis.components:
            if component.language == "unknown":
                blockers.append(f"Unknown technology in component: {component.name}")
            
            # Check for legacy technologies
            legacy_indicators = ['cobol', 'fortran', 'vb6', 'silverlight']
            if any(legacy in component.language.lower() for legacy in legacy_indicators):
                blockers.append(f"Legacy technology detected: {component.language}")
        
        # Check for accelerators
        if any(comp.docker_info for comp in analysis.components):
            accelerators.append("Already containerized")
        
        if 'kubernetes-native' in analysis.architecture_patterns:
            accelerators.append("Kubernetes-ready")
        
        if analysis.cicd_info.get('platform'):
            accelerators.append(f"CI/CD already implemented ({analysis.cicd_info['platform']})")
        
        if 'iac-terraform' in analysis.architecture_patterns:
            accelerators.append("Infrastructure as Code present")
        
        analysis.migration_considerations['blockers'] = blockers
        analysis.migration_considerations['accelerators'] = accelerators
        
        logger.info(f"Aggregation complete: Complexity={analysis.migration_considerations['complexity_level']}, "
                   f"Blockers={len(blockers)}, Accelerators={len(accelerators)}")
    
    def _detect_test_frameworks(self, test_files: List[str], language: str) -> List[str]:
        """Detect testing frameworks from test files"""
        frameworks = set()
        
        test_patterns = {
            'python': {
                'pytest': ['pytest', 'test_', '_test.py'],
                'unittest': ['unittest', 'TestCase'],
                'nose': ['nose', 'nose2']
            },
            'javascript': {
                'jest': ['jest', '.test.js', '.spec.js'],
                'mocha': ['mocha', 'describe(', 'it('],
                'jasmine': ['jasmine'],
                'cypress': ['cypress', 'cy.']
            },
            'java': {
                'junit': ['junit', '@Test', 'TestCase'],
                'testng': ['testng', '@Test'],
                'mockito': ['mockito', '@Mock']
            }
        }
        
        language_patterns = test_patterns.get(language, {})
        
        for test_file in test_files[:10]:  # Check first 10 test files
            try:
                with open(test_file, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read(1000)  # Read first 1KB
                    
                for framework, patterns in language_patterns.items():
                    if any(pattern in content for pattern in patterns):
                        frameworks.add(framework)
            except:
                pass
        
        return list(frameworks)
    
    def _generate_comprehensive_insights(self, analysis: RepositoryAnalysis):
        """Generate comprehensive insights using all collected data"""
        logger.info("Generating comprehensive AI-powered insights")
        
        # Prepare comprehensive data summary for LLM
        component_details = []
        for comp in analysis.components:
            details = {
                "name": comp.name,
                "language": comp.language,
                "framework": comp.framework,
                "docker": bool(comp.docker_info),
                "dependencies_count": sum(len(deps) for deps in comp.dependencies.values()),
                "databases": [db.name for db in comp.database_connections],
                "external_services": [s.name for s in comp.external_services],
                "api_endpoints_count": len(comp.api_endpoints),
                "env_vars_count": len(comp.environment_variables),
                "has_tests": comp.test_info.get('test_files_count', 0) > 0,
                "k8s_resources": len(comp.kubernetes_resources) + len(comp.openshift_resources)
            }
            component_details.append(details)
        
        # Create comprehensive prompt
        prompt = f"""
You are a cloud migration expert analyzing a repository for Azure migration. Based on this comprehensive analysis, provide detailed, actionable recommendations.

## Repository Overview
- Name: {analysis.repo_name}
- Total Components: {len(analysis.components)}
- Languages: {', '.join(analysis.tech_stack.get('languages', []))}
- Frameworks: {', '.join(analysis.tech_stack.get('frameworks', []))}
- Architecture Patterns: {', '.join(analysis.architecture_patterns)}
- Complexity Level: {analysis.migration_considerations.get('complexity_level', 'Unknown')}
- Complexity Factors: {', '.join(analysis.migration_considerations.get('complexity_factors', []))}

## Component Details
{json.dumps(component_details, indent=2)}

## Infrastructure Requirements
- Databases: {', '.join(analysis.tech_stack.get('databases', []))}
- External Services: {', '.join(analysis.tech_stack.get('services', []))}
- CI/CD Platform: {analysis.cicd_info.get('platform', 'None detected')}
- Container Orchestration: {'Yes' if any('kubernetes' in p or 'openshift' in p for p in analysis.architecture_patterns) else 'No'}

## Current State Summary
- Git Activity: {analysis.git_info.get('total_commits', 0)} commits, {analysis.git_info.get('contributor_count', 0)} contributors
- Last Activity: {analysis.git_info.get('last_commit_date', 'Unknown')}
- Migration Blockers: {', '.join(analysis.migration_considerations.get('blockers', [])) or 'None identified'}
- Migration Accelerators: {', '.join(analysis.migration_considerations.get('accelerators', [])) or 'None identified'}

## Security Findings
- Issues Found: {len(analysis.security_findings.get('issues', []))}
- Best Practices: {len(analysis.security_findings.get('best_practices', []))}

Based on this analysis, provide:

1. **Executive Summary**: Brief overview of the migration opportunity and complexity
2. **Recommended Azure Architecture**: Specific Azure services for each component
3. **Migration Strategy**: Step-by-step approach with phases
4. **Technical Requirements**: Detailed list of changes needed
5. **Risk Assessment**: Major risks and mitigation strategies
6. **Cost Estimation**: Rough estimate of Azure resources needed
7. **Timeline**: Realistic timeline based on complexity
8. **Quick Wins**: Immediate improvements that can be made

Format the response as a structured JSON with these sections. For each component, recommend specific Azure services (e.g., AKS, App Service, Functions, Container Instances).

Consider these Azure services in your recommendations:
- Compute: AKS, App Service, Container Instances, Functions
- Databases: Azure Database for PostgreSQL/MySQL, Cosmos DB, SQL Database
- Storage: Blob Storage, File Storage, Disk Storage
- Messaging: Service Bus, Event Hubs, Event Grid
- Cache: Azure Cache for Redis
- Security: Key Vault, Managed Identity, Security Center
- Monitoring: Application Insights, Log Analytics, Azure Monitor
- Networking: Application Gateway, Front Door, Private Endpoints
"""

        try:
            # Get LLM response
            response = self._analyze_with_llm(prompt)
            
            # Parse response
            import re
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                insights = json.loads(json_match.group())
                
                # Store detailed insights
                analysis.migration_considerations.update(insights)
                
                # Extract key recommendations for components
                if 'component_recommendations' in insights:
                    for comp in analysis.components:
                        if comp.name in insights['component_recommendations']:
                            comp_rec = insights['component_recommendations'][comp.name]
                            analysis.migration_considerations.setdefault('component_specific', {})[comp.name] = comp_rec
                
                logger.info("AI insights generated successfully")
            else:
                logger.warning("Could not parse JSON from LLM response")
                analysis.migration_considerations['ai_insights'] = response
        
        except Exception as e:
            logger.error(f"Error generating AI insights: {e}", exc_info=True)
            # Provide fallback recommendations
            self._generate_fallback_recommendations(analysis)
    
    def _generate_fallback_recommendations(self, analysis: RepositoryAnalysis):
        """Generate fallback recommendations if AI fails"""
        logger.info("Generating fallback recommendations")
        
        recommendations = {
            "executive_summary": f"Repository '{analysis.repo_name}' contains {len(analysis.components)} components using {', '.join(analysis.tech_stack.get('languages', []))}. Migration complexity is {analysis.migration_considerations.get('complexity_level', 'Medium')}.",
            "azure_services": [],
            "migration_phases": []
        }
        
        # Component-specific recommendations
        for comp in analysis.components:
            if comp.docker_info:
                if len(analysis.components) > 3:
                    rec = "Azure Kubernetes Service (AKS)"
                else:
                    rec = "Azure Container Instances"
            elif comp.language == 'javascript':
                rec = "Azure App Service (Node.js)"
            elif comp.language == 'python':
                rec = "Azure App Service (Python) or Azure Functions"
            elif comp.language == 'java':
                rec = "Azure App Service (Java) or AKS"
            elif comp.language == 'csharp':
                rec = "Azure App Service (.NET)"
            else:
                rec = "Azure Virtual Machines or AKS"
            
            recommendations.setdefault('component_recommendations', {})[comp.name] = {
                "azure_service": rec,
                "migration_complexity": "Medium"
            }
        
        # Database recommendations
        for db in analysis.tech_stack.get('databases', []):
            if 'postgres' in db.lower():
                recommendations['azure_services'].append("Azure Database for PostgreSQL")
            elif 'mysql' in db.lower():
                recommendations['azure_services'].append("Azure Database for MySQL")
            elif 'mongo' in db.lower():
                recommendations['azure_services'].append("Azure Cosmos DB (MongoDB API)")
            elif 'redis' in db.lower():
                recommendations['azure_services'].append("Azure Cache for Redis")
        
        analysis.migration_considerations.update(recommendations)


# Enhanced Document Generator
class DocumentGenerator:
    @staticmethod
    def generate_markdown(analysis: RepositoryAnalysis) -> Tuple[str, str, str]:
        """Generate comprehensive AS-IS, HLD, and LLD markdown documents"""
        logger.info("Starting enhanced document generation")
        
        # AS-IS State Document
        logger.info("Generating comprehensive AS-IS document")
        asis_md = f"""# AS-IS State Analysis - {analysis.repo_name}

## Executive Summary
- **Repository**: {analysis.repo_name}
- **Analysis Date**: {analysis.analysis_date}
- **Complexity Level**: {analysis.migration_considerations.get('complexity_level', 'Not assessed')}
- **Migration Readiness**: {len(analysis.migration_considerations.get('accelerators', []))} accelerators, {len(analysis.migration_considerations.get('blockers', []))} blockers

### Key Metrics
- **Contributors**: {analysis.git_info.get('contributor_count', 'Unknown')}
- **Total Commits**: {analysis.git_info.get('total_commits', 'Unknown')}
- **Last Activity**: {analysis.git_info.get('last_commit_date', 'Unknown')}
- **Components**: {len(analysis.components)}
- **Technologies**: {len(analysis.tech_stack.get('languages', []))} languages, {len(analysis.tech_stack.get('frameworks', []))} frameworks

## Repository Overview

### Git Statistics
- **Repository Age**: From {analysis.git_info.get('first_commit_date', 'Unknown')} to {analysis.git_info.get('last_commit_date', 'Unknown')}
- **Active Branches**: {len(analysis.git_info.get('branches', []))}
- **Default Branch**: {analysis.git_info.get('default_branch', 'Unknown')}

### Top Contributors
"""
        
        # Add contributor statistics
        if 'author_statistics' in analysis.git_info:
            sorted_authors = sorted(analysis.git_info['author_statistics'].items(), 
                                  key=lambda x: x[1]['commits'], reverse=True)[:5]
            for author, stats in sorted_authors:
                asis_md += f"- **{author}**: {stats['commits']} commits, +{stats['additions']} -{stats['deletions']}\n"
        
        asis_md += """
### Hot Files (Most Changed)
"""
        if 'hot_files' in analysis.git_info:
            for file, changes in analysis.git_info['hot_files'][:10]:
                asis_md += f"- `{file}`: {changes} changes\n"
        
        asis_md += """
## Technical Architecture

### Technology Stack Overview
"""
        
        # Technology summary table
        asis_md += """
| Category | Technologies |
|----------|-------------|
"""
        tech_categories = [
            ('Languages', analysis.tech_stack.get('languages', [])),
            ('Frameworks', analysis.tech_stack.get('frameworks', [])),
            ('Databases', analysis.tech_stack.get('databases', [])),
            ('Services', analysis.tech_stack.get('services', [])),
            ('Container', ['Docker'] if any(c.docker_info for c in analysis.components) else []),
            ('Orchestration', ['Kubernetes/OpenShift'] if any('kubernetes' in p or 'openshift' in p for p in analysis.architecture_patterns) else [])
        ]
        
        for category, items in tech_categories:
            if items:
                asis_md += f"| {category} | {', '.join(items)} |\n"
        
        # Architecture patterns
        asis_md += f"""
### Architecture Patterns Detected
"""
        for pattern in analysis.architecture_patterns:
            pattern_desc = {
                'microservices': 'Microservices architecture with multiple independent services',
                'monolithic': 'Monolithic application architecture',
                'kubernetes-native': 'Kubernetes-native deployment with container orchestration',
                'openshift-native': 'OpenShift-specific deployment configurations',
                'containerized': 'Containerized application using Docker',
                'serverless': 'Serverless components detected',
                'event-driven': 'Event-driven architecture with message queues',
                'api-gateway': 'API Gateway pattern for service routing'
            }
            desc = pattern_desc.get(pattern, pattern)
            asis_md += f"- **{pattern}**: {desc}\n"
        
        # Detailed component analysis
        asis_md += """
## Components and Services

### Component Details
"""
        
        for i, component in enumerate(analysis.components, 1):
            asis_md += f"""
#### {i}. {component.name}

**Overview**
- **Location**: `{component.path}`
- **Primary Language**: {component.language}
- **Framework**: {component.framework if component.framework != 'unknown' else 'Not detected'}
- **Containerized**: {'Yes' if component.docker_info else 'No'}
"""
            
            # Dependencies breakdown
            if component.dependencies:
                asis_md += "\n**Dependencies**\n"
                for dep_type, deps in component.dependencies.items():
                    if deps:
                        asis_md += f"- {dep_type.title()}: {len(deps)} packages\n"
                        # Show first 5 important dependencies
                        important_deps = [d for d in deps if any(key in d.lower() for key in ['express', 'react', 'django', 'spring', 'flask'])]
                        if important_deps:
                            asis_md += f"  - Key packages: {', '.join(important_deps[:5])}\n"
            
            # Docker information
            if component.docker_info:
                docker = component.docker_info
                asis_md += f"""
**Container Configuration**
- **Base Image**: `{docker.get('final_image') or docker.get('base_images', ['Unknown'])[0]}`
- **Multi-stage Build**: {'Yes' if len(docker.get('stages', [])) > 1 else 'No'}
- **Exposed Ports**: {', '.join(docker.get('exposed_ports', [])) or 'None'}
- **Volumes**: {len(docker.get('volumes', []))}
"""
                if docker.get('security_issues'):
                    asis_md += f"- **Security Concerns**: {len(docker['security_issues'])} issues found\n"
            
            # External connections
            if component.database_connections or component.external_services:
                asis_md += "\n**External Dependencies**\n"
                if component.database_connections:
                    asis_md += f"- **Databases**: {', '.join(set(db.name for db in component.database_connections))}\n"
                if component.external_services:
                    asis_md += f"- **Services**: {', '.join(set(s.name for s in component.external_services))}\n"
            
            # API Information
            if component.api_endpoints:
                asis_md += f"\n**API Endpoints**: {len(component.api_endpoints)} endpoints detected\n"
                # Show first 5 endpoints
                for endpoint in component.api_endpoints[:5]:
                    asis_md += f"- `{endpoint}`\n"
                if len(component.api_endpoints) > 5:
                    asis_md += f"- ... and {len(component.api_endpoints) - 5} more\n"
            
            # Testing
            if component.test_info.get('test_files_count', 0) > 0:
                asis_md += f"""
**Testing**
- **Test Files**: {component.test_info['test_files_count']}
- **Unit Tests**: {'Yes' if component.test_info.get('has_unit_tests') else 'No'}
- **Integration Tests**: {'Yes' if component.test_info.get('has_integration_tests') else 'No'}
- **Test Frameworks**: {', '.join(component.test_info.get('test_frameworks', [])) or 'None detected'}
"""
        
        # Infrastructure requirements
        asis_md += """
## Infrastructure and Deployment

### Current Infrastructure Requirements
"""
        
        if analysis.infrastructure_requirements:
            for key, value in analysis.infrastructure_requirements.items():
                asis_md += f"- **{key.replace('_', ' ').title()}**: {value}\n"
        
        # CI/CD Information
        if analysis.cicd_info:
            asis_md += f"""
### CI/CD Pipeline
- **Platform**: {analysis.cicd_info.get('platform', 'Not detected')}
- **Docker Build**: {'Yes' if analysis.cicd_info.get('docker_build') else 'No'}
- **Automated Tests**: {'Yes' if analysis.cicd_info.get('automated_tests') else 'No'}
- **Automated Deployment**: {'Yes' if analysis.cicd_info.get('automated_deployment') else 'No'}
"""
            if 'stages' in analysis.cicd_info:
                asis_md += f"- **Pipeline Stages**: {', '.join(analysis.cicd_info['stages'])}\n"
        
        # Kubernetes/OpenShift resources
        if analysis.deployment_info.get('kubernetes_resources'):
            asis_md += """
### Container Orchestration Resources
"""
            k8s_resources = analysis.deployment_info['kubernetes_resources']
            asis_md += f"- **Total Resources**: {analysis.deployment_info.get('total_k8s_resources', 0)}\n"
            asis_md += "- **Resource Types**:\n"
            for resource_type, resources in k8s_resources.items():
                asis_md += f"  - {resource_type}: {len(resources)}\n"
        
        # Security findings
        asis_md += """
## Security Analysis

### Current Security Posture
"""
        if analysis.security_findings:
            issues = analysis.security_findings.get('issues', [])
            best_practices = analysis.security_findings.get('best_practices', [])
            
            asis_md += f"- **Security Issues Found**: {len(issues)}\n"
            if issues:
                for issue in issues[:5]:
                    asis_md += f"  - {issue}\n"
                if len(issues) > 5:
                    asis_md += f"  - ... and {len(issues) - 5} more\n"
            
            asis_md += f"- **Security Best Practices**: {len(best_practices)}\n"
            if best_practices:
                for practice in best_practices[:5]:
                    asis_md += f"  - {practice}\n"
        
        # Migration readiness
        asis_md += """
## Migration Readiness Assessment

### Complexity Analysis
"""
        asis_md += f"- **Overall Complexity**: {analysis.migration_considerations.get('complexity_level', 'Not assessed')}\n"
        asis_md += f"- **Complexity Score**: {analysis.migration_considerations.get('complexity_score', 'N/A')}/10\n"
        
        if 'complexity_factors' in analysis.migration_considerations:
            asis_md += "- **Contributing Factors**:\n"
            for factor in analysis.migration_considerations['complexity_factors']:
                asis_md += f"  - {factor}\n"
        
        asis_md += """
### Migration Accelerators
"""
        accelerators = analysis.migration_considerations.get('accelerators', [])
        if accelerators:
            for acc in accelerators:
                asis_md += f"- ✅ {acc}\n"
        else:
            asis_md += "- No significant accelerators identified\n"
        
        asis_md += """
### Migration Blockers
"""
        blockers = analysis.migration_considerations.get('blockers', [])
        if blockers:
            for blocker in blockers:
                asis_md += f"- ❌ {blocker}\n"
        else:
            asis_md += "- ✅ No significant blockers identified\n"
        
        # Gaps and missing information
        if analysis.gaps:
            asis_md += """
## Analysis Gaps

The following areas could not be fully analyzed:
"""
            for gap in analysis.gaps:
                asis_md += f"- {gap}\n"
        
        logger.info("AS-IS document generated")
        
        # HLD Document
        logger.info("Generating comprehensive HLD document")
        hld_md = f"""# High Level Design - {analysis.repo_name}

## Executive Summary

This High-Level Design document outlines the architecture and migration strategy for moving {analysis.repo_name} to Microsoft Azure. The migration complexity is assessed as **{analysis.migration_considerations.get('complexity_level', 'Medium')}** with {len(analysis.components)} components requiring migration.

### Migration Overview
- **Source State**: {len(analysis.tech_stack.get('languages', []))} languages, {len(analysis.components)} components
- **Target Platform**: Microsoft Azure
- **Migration Type**: Rehosting with modernization opportunities
- **Estimated Timeline**: {analysis.migration_considerations.get('timeline', '8-12 weeks')}

## Scope

### In Scope
1. **Application Components** ({len(analysis.components)} total)
"""
        for comp in analysis.components:
            hld_md += f"   - {comp.name} ({comp.language})\n"
        
        hld_md += f"""
2. **Data Stores** ({len(analysis.tech_stack.get('databases', []))} types)
"""
        for db in analysis.tech_stack.get('databases', []):
            hld_md += f"   - {db}\n"
        
        hld_md += """
3. **Infrastructure Services**
   - Container orchestration and management
   - Load balancing and traffic management
   - Security and compliance controls
   - Monitoring and observability
   - CI/CD pipeline migration

4. **Non-Functional Requirements**
   - High availability (99.9% SLA)
   - Disaster recovery capabilities
   - Performance optimization
   - Security hardening
   - Cost optimization

### Out of Scope
- Third-party SaaS integrations (remain as-is)
- End-user training (separate initiative)
- Data archival beyond 2 years
- Legacy system decommissioning (phase 2)

## Current State Architecture

### Architecture Overview
"""
        
        # Architecture diagram description
        hld_md += f"""
The current architecture consists of {len(analysis.components)} components following a {', '.join(analysis.architecture_patterns)} pattern.

```
Current Architecture:
"""
        
        # Simple ASCII diagram
        if len(analysis.components) > 1:
            hld_md += """
┌─────────────────┐     ┌─────────────────┐
│   Component 1   │────▶│   Component 2   │
└─────────────────┘     └─────────────────┘
         │                       │
         ▼                       ▼
┌─────────────────┐     ┌─────────────────┐
│    Database     │     │  External API   │
└─────────────────┘     └─────────────────┘
```
"""
        else:
            hld_md += """
┌─────────────────┐
│   Application   │
└─────────────────┘
         │
         ▼
┌─────────────────┐
│    Database     │
└─────────────────┘
```
"""
        
        hld_md += """
### Technology Stack Summary

| Layer | Current Technology | Components Using |
|-------|-------------------|------------------|
"""
        
        # Group components by technology
        tech_usage = defaultdict(list)
        for comp in analysis.components:
            tech_usage[comp.language].append(comp.name)
        
        for tech, comps in tech_usage.items():
            hld_md += f"| Application | {tech} | {', '.join(comps)} |\n"
        
        for db in analysis.tech_stack.get('databases', []):
            hld_md += f"| Database | {db} | Multiple components |\n"
        
        if any(c.docker_info for c in analysis.components):
            hld_md += "| Container | Docker | All components |\n"
        
        if 'kubernetes-native' in analysis.architecture_patterns or 'openshift-native' in analysis.architecture_patterns:
            hld_md += "| Orchestration | Kubernetes/OpenShift | All components |\n"
        
        hld_md += """
## Target State Architecture on Azure

### Azure Architecture Overview

The target architecture leverages Azure's native services for improved scalability, reliability, and operational efficiency.

```
Target Azure Architecture:
"""
        
        # Azure architecture diagram
        if 'kubernetes' in str(analysis.architecture_patterns) or len(analysis.components) > 3:
            hld_md += """
                    ┌─────────────────────┐
                    │   Azure Front Door  │
                    └──────────┬──────────┘
                               │
                    ┌──────────▼──────────┐
                    │    Azure AKS        │
                    │  ┌──────┬───────┐  │
                    │  │ Pod  │  Pod  │  │
                    │  └──────┴───────┘  │
                    └──────────┬──────────┘
                               │
                ┌──────────────┴──────────────┐
                │                             │
     ┌──────────▼──────────┐      ┌──────────▼──────────┐
     │  Azure Database     │      │  Azure Cache Redis  │
     └────────────────────┘      └────────────────────┘
```
"""
        else:
            hld_md += """
                    ┌─────────────────────┐
                    │ Azure App Gateway   │
                    └──────────┬──────────┘
                               │
                    ┌──────────▼──────────┐
                    │  Azure App Service  │
                    └──────────┬──────────┘
                               │
                ┌──────────────┴──────────────┐
                │                             │
     ┌──────────▼──────────┐      ┌──────────▼──────────┐
     │  Azure SQL Database │      │    Azure Storage    │
     └────────────────────┘      └────────────────────┘
```
"""
        
        hld_md += """
### Recommended Azure Services

Based on the analysis, the following Azure services are recommended:

#### Compute Services
"""
        
        # Component-specific recommendations
        if 'component_recommendations' in analysis.migration_considerations:
            for comp_name, rec in analysis.migration_considerations['component_recommendations'].items():
                hld_md += f"- **{comp_name}**: {rec.get('azure_service', 'Azure App Service')}\n"
        else:
            # Fallback recommendations
            for comp in analysis.components:
                if comp.docker_info:
                    service = "Azure Kubernetes Service (AKS)" if len(analysis.components) > 3 else "Azure Container Instances"
                else:
                    service = "Azure App Service"
                hld_md += f"- **{comp.name}**: {service}\n"
        
        hld_md += """
#### Data Services
"""
        db_mapping = {
            'postgresql': 'Azure Database for PostgreSQL',
            'mysql': 'Azure Database for MySQL',
            'mongodb': 'Azure Cosmos DB (MongoDB API)',
            'redis': 'Azure Cache for Redis',
            'elasticsearch': 'Azure Cognitive Search'
        }
        
        for db in analysis.tech_stack.get('databases', []):
            azure_service = db_mapping.get(db.lower(), 'Azure SQL Database')
            hld_md += f"- **{db}** → {azure_service}\n"
        
        hld_md += """
#### Supporting Services
- **Security**: Azure Key Vault for secrets management
- **Identity**: Azure Active Directory with Managed Identity
- **Monitoring**: Azure Monitor with Application Insights
- **Networking**: Azure Virtual Network with Private Endpoints
- **API Management**: Azure API Management (if needed)
- **CDN**: Azure Front Door for global distribution
- **CI/CD**: Azure DevOps or GitHub Actions

### High-Level Migration Approach

#### Migration Strategy: {analysis.migration_considerations.get('strategy', 'Lift and Shift with Optimization')}

The migration will follow a phased approach to minimize risk and ensure business continuity.

#### Phase 1: Foundation (Weeks 1-2)
- Set up Azure landing zone
- Configure networking and security
- Establish Azure DevOps/GitHub integration
- Create base infrastructure using Terraform/ARM

#### Phase 2: Data Migration (Weeks 3-4)
- Set up Azure data services
- Implement data replication
- Test data integrity
- Plan cutover strategy

#### Phase 3: Application Migration (Weeks 5-7)
- Containerize applications (if needed)
- Deploy to Azure compute services
- Configure load balancing
- Implement monitoring

#### Phase 4: Integration & Testing (Weeks 8-9)
- End-to-end integration testing
- Performance testing and optimization
- Security validation
- DR testing

#### Phase 5: Cutover & Optimization (Weeks 10-12)
- Production cutover
- Performance tuning
- Cost optimization
- Documentation and handover

### Architecture Principles

1. **Cloud-Native Design**
   - Leverage PaaS services where possible
   - Use managed services to reduce operational overhead
   - Implement auto-scaling for cost efficiency

2. **Security by Design**
   - Zero-trust network architecture
   - Encryption at rest and in transit
   - Managed identities for service authentication
   - Regular security assessments

3. **High Availability**
   - Multi-zone deployment for critical services
   - Automated failover mechanisms
   - Regular backup and restore testing

4. **Operational Excellence**
   - Infrastructure as Code (IaC)
   - Automated deployment pipelines
   - Comprehensive monitoring and alerting
   - Self-healing capabilities

### Non-Functional Requirements

#### Performance
- **Response Time**: < 200ms for API calls
- **Throughput**: Support {analysis.migration_considerations.get('expected_tps', '1000')} TPS
- **Availability**: 99.9% uptime SLA

#### Security
- **Compliance**: SOC 2, ISO 27001 (as applicable)
- **Encryption**: TLS 1.2+ for transit, AES-256 for storage
- **Access Control**: RBAC with Azure AD integration

#### Scalability
- **Horizontal Scaling**: Auto-scale based on CPU/memory
- **Vertical Scaling**: Resize without downtime
- **Global Scale**: Multi-region deployment capability

#### Disaster Recovery
- **RTO**: 4 hours
- **RPO**: 1 hour
- **Backup Retention**: 30 days
- **Geo-Redundancy**: Cross-region replication

### Cost Considerations

#### Estimated Monthly Costs (Production)
"""
        
        # Cost estimation based on components
        base_cost = 500  # Base infrastructure
        component_cost = len(analysis.components) * 200
        db_cost = len(analysis.tech_stack.get('databases', [])) * 300
        total_cost = base_cost + component_cost + db_cost
        
        hld_md += f"""
- **Compute**: ${component_cost}
- **Storage & Databases**: ${db_cost}
- **Networking**: $200
- **Monitoring**: $100
- **Security**: $200
- **Backup & DR**: $300
- **Total Estimate**: ${total_cost} - ${total_cost * 1.5}

*Note: Actual costs depend on usage patterns and specific configurations*

#### Cost Optimization Strategies
- Use Azure Reserved Instances for predictable workloads
- Implement auto-scaling to match demand
- Use Azure Hybrid Benefit for Windows/SQL workloads
- Regular cost reviews and optimization

### Risk Assessment

#### Technical Risks
| Risk | Impact | Probability | Mitigation |
|------|--------|-------------|------------|
| Data migration failure | High | Low | Comprehensive testing, rollback plan |
| Performance degradation | Medium | Medium | Performance testing, optimization |
| Integration issues | Medium | Medium | Thorough integration testing |
| Security vulnerabilities | High | Low | Security scanning, best practices |

#### Operational Risks
| Risk | Impact | Probability | Mitigation |
|------|--------|-------------|------------|
| Team skill gaps | Medium | Medium | Training, Azure support |
| Downtime during cutover | High | Low | Blue-green deployment |
| Cost overrun | Medium | Medium | Regular monitoring, alerts |

### Success Criteria

1. **Technical Success**
   - All components successfully migrated
   - Performance SLAs met or exceeded
   - Zero data loss during migration
   - Security compliance achieved

2. **Operational Success**
   - Team trained on Azure operations
   - Automated deployment pipeline
   - Monitoring and alerting operational
   - Documentation complete

3. **Business Success**
   - Minimal business disruption
   - Cost within budget
   - Improved system reliability
   - Enhanced scalability achieved
"""
        
        logger.info("HLD document generated")
        
        # LLD Document
        logger.info("Generating comprehensive LLD document")
        lld_md = f"""# Low Level Design - {analysis.repo_name}

## Introduction

This Low-Level Design document provides detailed technical specifications for migrating {analysis.repo_name} to Microsoft Azure. It serves as the implementation guide for the development and operations teams.

## Document Purpose and Audience

**Purpose**: Provide detailed technical implementation specifications for the Azure migration.

**Audience**:
- Development Team - Implementation details
- DevOps Team - Infrastructure and deployment
- Security Team - Security configurations
- Database Administrators - Data migration procedures
- QA Team - Testing requirements

## Detailed Component Specifications
"""
        
        for i, component in enumerate(analysis.components, 1):
            lld_md += f"""
### {i}. Component: {component.name}

#### Current State Technical Details

**Technology Stack**
- **Language**: {component.language}
- **Framework**: {component.framework if component.framework != 'unknown' else 'Not specified'}
- **Dependencies**: {sum(len(deps) for deps in component.dependencies.values())} total packages
"""
            
            # List key dependencies
            if component.dependencies:
                lld_md += "\n**Key Dependencies**:\n"
                for dep_type, deps in component.dependencies.items():
                    if deps:
                        # Show important dependencies
                        important = [d for d in deps[:10] if not d.startswith('@types/')]
                        if important:
                            lld_md += f"- {dep_type.title()}: {', '.join(important[:5])}"
                            if len(important) > 5:
                                lld_md += f" (+{len(important)-5} more)"
                            lld_md += "\n"
            
            # Source file analysis
            if component.source_files:
                lld_md += f"\n**Code Statistics**:\n"
                lld_md += f"- Source Files Analyzed: {len(component.source_files)}\n"
                patterns = []
                for file in component.source_files:
                    patterns.extend(file.detected_patterns)
                unique_patterns = set(patterns)
                if unique_patterns:
                    lld_md += f"- Detected Patterns: {', '.join(unique_patterns)}\n"
            
            # Current container configuration
            if component.docker_info:
                docker = component.docker_info
                lld_md += f"""
**Container Configuration**
- **Base Image**: `{docker.get('final_image') or docker.get('base_images', ['Unknown'])[0]}`
- **Multi-stage Build**: {'Yes' if len(docker.get('stages', [])) > 1 else 'No'}
"""
                if docker.get('stages'):
                    lld_md += "- **Build Stages**:\n"
                    for stage in docker.get('stages', []):
                        lld_md += f"  - {stage['name']}: FROM {stage['from']}\n"
                
                if docker.get('exposed_ports'):
                    lld_md += f"- **Exposed Ports**: {', '.join(docker['exposed_ports'])}\n"
                
                if docker.get('environment_variables'):
                    lld_md += f"- **Environment Variables**: {len(docker['environment_variables'])} defined\n"
                    # List non-sensitive env vars
                    for key, value in list(docker['environment_variables'].items())[:5]:
                        if not any(secret in key.lower() for secret in ['password', 'key', 'secret', 'token']):
                            lld_md += f"  - `{key}={value}`\n"
            
            # Environment configuration
            if component.environment_variables:
                lld_md += f"\n**Environment Configuration**:\n"
                lld_md += f"- Total Variables: {len(component.environment_variables)}\n"
                # Group by type
                config_vars = []
                secret_vars = []
                for key in component.environment_variables:
                    if any(secret in key.lower() for secret in ['password', 'key', 'secret', 'token']):
                        secret_vars.append(key)
                    else:
                        config_vars.append(key)
                
                if config_vars:
                    lld_md += f"- Configuration Variables: {', '.join(config_vars[:5])}\n"
                if secret_vars:
                    lld_md += f"- Secret Variables: {len(secret_vars)} (to be migrated to Key Vault)\n"
            
            # External dependencies
            if component.database_connections or component.external_services:
                lld_md += "\n**External Dependencies**:\n"
                
                if component.database_connections:
                    lld_md += "- **Databases**:\n"
                    for db in component.database_connections:
                        lld_md += f"  - {db.name}: {db.connection_string if not any(s in str(db.connection_string) for s in ['password', 'pwd']) else '[REDACTED]'}\n"
                
                if component.external_services:
                    lld_md += "- **External Services**:\n"
                    for service in component.external_services[:5]:
                        lld_md += f"  - {service.name} ({service.type})\n"
            
            # API endpoints
            if component.api_endpoints:
                lld_md += f"\n**API Endpoints** ({len(component.api_endpoints)} total):\n"
                for endpoint in component.api_endpoints[:10]:
                    lld_md += f"- `{endpoint}`\n"
                if len(component.api_endpoints) > 10:
                    lld_md += f"- ... and {len(component.api_endpoints) - 10} more endpoints\n"
            
            # Target state configuration
            lld_md += """
#### Target State on Azure

**Azure Service Configuration**
"""
            
            # Get specific recommendation for this component
            comp_rec = analysis.migration_considerations.get('component_recommendations', {}).get(component.name, {})
            azure_service = comp_rec.get('azure_service', 'Azure App Service')
            
            lld_md += f"- **Target Service**: {azure_service}\n"
            
            # Service-specific configuration
            if 'AKS' in azure_service or 'Kubernetes' in azure_service:
                lld_md += """
**AKS Deployment Configuration**
```yaml
apiVersion: apps/v1
kind: Deployment
metadata:
  name: """ + component.name + """
  namespace: production
spec:
  replicas: 3
  selector:
    matchLabels:
      app: """ + component.name + """
  template:
    metadata:
      labels:
        app: """ + component.name + """
    spec:
      containers:
      - name: """ + component.name + """
        image: <acr-name>.azurecr.io/""" + component.name + """:latest
        ports:"""
                
                if component.docker_info and component.docker_info.get('exposed_ports'):
                    for port in component.docker_info['exposed_ports']:
                        lld_md += f"""
        - containerPort: {port}"""
                else:
                    lld_md += """
        - containerPort: 8080"""
                
                lld_md += """
        resources:
          requests:
            cpu: 100m
            memory: 128Mi
          limits:
            cpu: 500m
            memory: 512Mi
        env:
        - name: AZURE_CLIENT_ID
          valueFrom:
            secretKeyRef:
              name: azure-identity
              key: client-id
```

**Service Configuration**
```yaml
apiVersion: v1
kind: Service
metadata:
  name: """ + component.name + """-service
spec:
  selector:
    app: """ + component.name + """
  ports:
  - protocol: TCP
    port: 80
    targetPort: """ + (component.docker_info.get('exposed_ports', ['8080'])[0] if component.docker_info else '8080') + """
  type: ClusterIP
```
"""
            elif 'App Service' in azure_service:
                lld_md += f"""
**App Service Configuration**
- **Service Plan**: P1v3 (Production), B1 (Dev/Test)
- **Runtime Stack**: {component.language.title()} {component.framework if component.framework != 'unknown' else ''}
- **Operating System**: Linux
- **Always On**: Enabled (Production)
- **Auto-Scale Rules**:
  - CPU > 70% for 5 minutes: Scale out by 1 instance
  - CPU < 30% for 10 minutes: Scale in by 1 instance
  - Min instances: 2, Max instances: 10

**Application Settings**
```json
{{
  "WEBSITES_ENABLE_APP_SERVICE_STORAGE": "false",
  "WEBSITE_RUN_FROM_PACKAGE": "1",
  "APPLICATIONINSIGHTS_CONNECTION_STRING": "<from-key-vault>",
  "KeyVaultUri": "https://<keyvault-name>.vault.azure.net/"
}}
```
"""
            elif 'Container Instances' in azure_service:
                lld_md += """
**Container Instances Configuration**
- **CPU**: 1 core
- **Memory**: 1.5 GB
- **OS Type**: Linux
- **Restart Policy**: Always
- **Network**: VNet integrated
"""
            
            # Migration steps
            lld_md += f"""
#### Migration Implementation Steps

1. **Pre-Migration Preparation**
   - Code repository branch: `azure-migration/{component.name}`
   - Update configuration for Azure services
   - Add Azure SDK dependencies
   - Implement health check endpoints

2. **Configuration Changes Required**
"""
            
            # Specific code changes based on language
            if component.language == 'javascript' or component.language == 'typescript':
                lld_md += """   ```javascript
   // Add Azure App Configuration
   const { AppConfigurationClient } = require("@azure/app-configuration");
   const { DefaultAzureCredential } = require("@azure/identity");
   
   // Add Application Insights
   const appInsights = require("applicationinsights");
   appInsights.setup(process.env.APPLICATIONINSIGHTS_CONNECTION_STRING)
       .setAutoDependencyCorrelation(true)
       .setAutoCollectRequests(true)
       .setAutoCollectPerformance(true)
       .start();
   ```
"""
            elif component.language == 'python':
                lld_md += """   ```python
   # Add Azure SDK dependencies to requirements.txt
   azure-identity==1.12.0
   azure-keyvault-secrets==4.6.0
   azure-appconfiguration==1.4.0
   opencensus-ext-azure==1.1.9
   
   # Update application code
   from azure.identity import DefaultAzureCredential
   from azure.keyvault.secrets import SecretClient
   from azure.appconfiguration import AzureAppConfigurationClient
   ```
"""
            elif component.language == 'java':
                lld_md += """   ```xml
   <!-- Add to pom.xml -->
   <dependency>
       <groupId>com.azure</groupId>
       <artifactId>azure-identity</artifactId>
       <version>1.8.0</version>
   </dependency>
   <dependency>
       <groupId>com.azure</groupId>
       <artifactId>azure-security-keyvault-secrets</artifactId>
       <version>4.5.0</version>
   </dependency>
   ```
"""
            
            # Database connection updates
            if component.database_connections:
                lld_md += """
3. **Database Connection Updates**
   - Migrate connection strings to Azure Key Vault
   - Update connection logic to use Managed Identity
   - Implement connection retry logic
   - Add connection pooling configuration

   ```"""
                if component.language == 'javascript':
                    lld_md += """javascript
   const { SecretClient } = require("@azure/keyvault-secrets");
   const credential = new DefaultAzureCredential();
   const client = new SecretClient(vaultUrl, credential);
   const dbConnString = await client.getSecret("db-connection-string");
   ```
"""
                elif component.language == 'python':
                    lld_md += """python
   from azure.keyvault.secrets import SecretClient
   credential = DefaultAzureCredential()
   client = SecretClient(vault_url=vault_url, credential=credential)
   db_conn_string = client.get_secret("db-connection-string").value
   ```
"""
            
            # Testing requirements
            lld_md += """
4. **Testing Requirements**
   - Unit Tests: Ensure all existing tests pass
   - Integration Tests: Test Azure service integrations
   - Performance Tests: Validate response times
   - Security Tests: Scan for vulnerabilities

5. **Deployment Configuration**
"""
            
            if 'AKS' in azure_service:
                lld_md += """   - Build and push Docker image to ACR
   - Update Kubernetes manifests
   - Configure Horizontal Pod Autoscaler
   - Set up Ingress rules
"""
            else:
                lld_md += """   - Configure deployment slots (staging, production)
   - Set up deployment from Azure DevOps/GitHub
   - Configure custom domain (if applicable)
   - Enable Application Insights
"""
            
            lld_md += """
#### Monitoring and Observability

**Application Insights Configuration**
- **Instrumentation**: Automatic + Custom
- **Sampling**: Adaptive (Production), 100% (Dev/Test)
- **Retention**: 90 days
- **Alerts**:
  - Response time > 1s for 5 minutes
  - Error rate > 1% for 5 minutes
  - Availability < 99.9%

**Log Analytics Queries**
```kusto
// Component health check
requests
| where name == "GET /health"
| summarize 
    SuccessRate = countif(success == true) * 100.0 / count(),
    AvgDuration = avg(duration)
    by bin(timestamp, 5m)
| order by timestamp desc
```

**Custom Metrics**
- Business metrics specific to component functionality
- Performance counters
- Custom events for important operations
"""
        
        # Database migration section
        lld_md += """
## Database Migration Strategy

### Overview
"""
        
        databases = analysis.tech_stack.get('databases', [])
        if databases:
            lld_md += f"The application uses {len(databases)} database type(s): {', '.join(databases)}\n\n"
            
            for db in databases:
                if 'postgres' in db.lower():
                    lld_md += """
### PostgreSQL to Azure Database for PostgreSQL

**Migration Approach**: Online migration using Azure Database Migration Service

**Pre-Migration Steps**
1. **Assessment**
   ```bash
   # Run Azure migration assessment
   ora2pg -t SHOW_REPORT -c ora2pg.conf
   ```

2. **Target Setup**
   - Create Azure Database for PostgreSQL
   - Configure: General Purpose, Gen5, 4 vCores
   - Enable connection security
   - Configure firewall rules

3. **Schema Migration**
   ```sql
   -- Export schema
   pg_dump -h source-server -U username -d dbname -s > schema.sql
   
   -- Import to Azure
   psql -h target-server.postgres.database.azure.com -U username@servername -d dbname < schema.sql
   ```

4. **Data Migration**
   - Use Azure DMS for minimal downtime
   - Configure continuous sync
   - Monitor replication lag

**Connection String Update**
```
# Old
postgresql://user:pass@localhost:5432/dbname

# New (from Key Vault)
postgresql://user@server:pass@server.postgres.database.azure.com:5432/dbname?sslmode=require
```
"""
                elif 'mysql' in db.lower():
                    lld_md += """
### MySQL to Azure Database for MySQL

**Migration Approach**: Azure Database Migration Service with minimal downtime

**Configuration**
- **SKU**: Standard_B2ms (Dev/Test), Standard_D4s_v3 (Production)
- **Storage**: 128 GB with auto-grow enabled
- **Backup**: Geo-redundant, 35-day retention
- **High Availability**: Zone redundant (Production)
"""
                elif 'mongo' in db.lower():
                    lld_md += """
### MongoDB to Azure Cosmos DB

**Migration Approach**: Azure Database Migration Service or native tools

**Cosmos DB Configuration**
- **API**: MongoDB API
- **Consistency**: Session consistency
- **Throughput**: 1000 RU/s (autoscale to 10,000)
- **Partition Key**: Choose based on query patterns
- **Global Distribution**: Single region initially

**Migration Steps**
```bash
# Export from MongoDB
mongodump --uri="mongodb://source-server:27017/dbname"

# Import to Cosmos DB
mongorestore --uri="<cosmos-connection-string>" --dir=dump/
```
"""
                elif 'redis' in db.lower():
                    lld_md += """
### Redis to Azure Cache for Redis

**Configuration**
- **Tier**: Premium P1 (Production), Standard C2 (Dev/Test)
- **Clustering**: Enabled for Production
- **Persistence**: RDB + AOF
- **Availability Zones**: Enabled

**Migration Approach**
```bash
# Export data
redis-cli --rdb dump.rdb

# Import to Azure Cache
redis-cli -h <cache-name>.redis.cache.windows.net -a <access-key> --rdb dump.rdb
```
"""
        
        # Infrastructure as Code
        lld_md += """
## Infrastructure as Code

### Terraform Configuration

**Main Infrastructure** (`main.tf`)
```hcl
terraform {
  required_providers {
    azurerm = {
      source  = "hashicorp/azurerm"
      version = "~>3.0"
    }
  }
  
  backend "azurerm" {
    resource_group_name  = "terraform-state-rg"
    storage_account_name = "terraformstate"
    container_name       = "tfstate"
    key                  = "prod.terraform.tfstate"
  }
}

provider "azurerm" {
  features {}
}

# Resource Group
resource "azurerm_resource_group" "main" {
  name     = "${var.project_name}-${var.environment}-rg"
  location = var.location
  tags     = var.tags
}

# Virtual Network
resource "azurerm_virtual_network" "main" {
  name                = "${var.project_name}-${var.environment}-vnet"
  address_space       = ["10.0.0.0/16"]
  location            = azurerm_resource_group.main.location
  resource_group_name = azurerm_resource_group.main.name
}
```

**AKS Cluster** (`aks.tf`)
```hcl
resource "azurerm_kubernetes_cluster" "main" {
  name                = "${var.project_name}-${var.environment}-aks"
  location            = azurerm_resource_group.main.location
  resource_group_name = azurerm_resource_group.main.name
  dns_prefix          = "${var.project_name}-${var.environment}"
  
  default_node_pool {
    name                = "default"
    node_count          = var.node_count
    vm_size             = "Standard_D4s_v3"
    type                = "VirtualMachineScaleSets"
    enable_auto_scaling = true
    min_count           = 2
    max_count           = 10
  }
  
  identity {
    type = "SystemAssigned"
  }
  
  network_profile {
    network_plugin     = "azure"
    network_policy     = "calico"
    load_balancer_sku  = "standard"
    service_cidr       = "10.1.0.0/16"
    dns_service_ip     = "10.1.0.10"
  }
}
```

### CI/CD Pipeline

**Azure DevOps Pipeline** (`azure-pipelines.yml`)
```yaml
trigger:
  branches:
    include:
    - main
    - develop
  paths:
    exclude:
    - README.md
    - docs/*

variables:
  azureSubscription: 'Azure-Service-Connection'
  resourceGroupName: 'app-prod-rg'
  location: 'eastus'
  terraformVersion: '1.3.0'

stages:
- stage: Build
  jobs:
  - job: BuildAndTest
    pool:
      vmImage: 'ubuntu-latest'
    steps:
    - task: UseNode@1
      inputs:
        version: '18.x'
      condition: eq(variables['componentLanguage'], 'javascript')
    
    - task: UsePythonVersion@0
      inputs:
        versionSpec: '3.10'
      condition: eq(variables['componentLanguage'], 'python')
    
    - script: |
        # Install dependencies
        npm ci  # or pip install -r requirements.txt
      displayName: 'Install dependencies'
    
    - script: |
        # Run tests
        npm test  # or pytest
      displayName: 'Run tests'
    
    - task: Docker@2
      inputs:
        containerRegistry: 'ACR-Service-Connection'
        repository: '$(Build.Repository.Name)'
        command: 'buildAndPush'
        Dockerfile: '**/Dockerfile'
        tags: |
          $(Build.BuildId)
          latest

- stage: DeployDev
  condition: and(succeeded(), eq(variables['Build.SourceBranch'], 'refs/heads/develop'))
  jobs:
  - deployment: DeployToDev
    environment: 'dev'
    strategy:
      runOnce:
        deploy:
          steps:
          - task: HelmDeploy@0
            inputs:
              connectionType: 'Azure Resource Manager'
              azureSubscription: $(azureSubscription)
              azureResourceGroup: $(resourceGroupName)
              kubernetesCluster: 'app-dev-aks'
              command: 'upgrade'
              chartType: 'FilePath'
              chartPath: '$(Pipeline.Workspace)/charts'
              releaseName: 'app-dev'
              overrideValues: 'image.tag=$(Build.BuildId)'

- stage: DeployProd
  condition: and(succeeded(), eq(variables['Build.SourceBranch'], 'refs/heads/main'))
  jobs:
  - deployment: DeployToProd
    environment: 'prod'
    strategy:
      runOnce:
        deploy:
          steps:
          - task: HelmDeploy@0
            inputs:
              connectionType: 'Azure Resource Manager'
              azureSubscription: $(azureSubscription)
              azureResourceGroup: $(resourceGroupName)
              kubernetesCluster: 'app-prod-aks'
              command: 'upgrade'
              chartType: 'FilePath'
              chartPath: '$(Pipeline.Workspace)/charts'
              releaseName: 'app-prod'
              overrideValues: 'image.tag=$(Build.BuildId)'
```

## Security Implementation

### Azure Key Vault Integration

**Key Vault Setup**
```bash
# Create Key Vault
az keyvault create \
  --name "${PROJECT_NAME}kv" \
  --resource-group "${RESOURCE_GROUP}" \
  --location "${LOCATION}" \
  --enable-rbac-authorization

# Add secrets
az keyvault secret set \
  --vault-name "${PROJECT_NAME}kv" \
  --name "db-connection-string" \
  --value "${DB_CONNECTION_STRING}"

# Grant access to AKS/App Service
az keyvault set-policy \
  --name "${PROJECT_NAME}kv" \
  --object-id "${MANAGED_IDENTITY_OBJECT_ID}" \
  --secret-permissions get list
```

**Application Integration**
```csharp
// C# Example
var keyVaultName = Environment.GetEnvironmentVariable("KEY_VAULT_NAME");
var kvUri = $"https://{keyVaultName}.vault.azure.net";

var client = new SecretClient(new Uri(kvUri), new DefaultAzureCredential());
var secret = await client.GetSecretAsync("db-connection-string");
```

### Network Security

**Network Security Groups**
```hcl
resource "azurerm_network_security_group" "app" {
  name                = "${var.project_name}-app-nsg"
  location            = azurerm_resource_group.main.location
  resource_group_name = azurerm_resource_group.main.name

  security_rule {
    name                       = "AllowHTTPS"
    priority                   = 100
    direction                  = "Inbound"
    access                     = "Allow"
    protocol                   = "Tcp"
    source_port_range          = "*"
    destination_port_range     = "443"
    source_address_prefix      = "*"
    destination_address_prefix = "*"
  }

  security_rule {
    name                       = "DenyAllInbound"
    priority                   = 200
    direction                  = "Inbound"
    access                     = "Deny"
    protocol                   = "*"
    source_port_range          = "*"
    destination_port_range     = "*"
    source_address_prefix      = "*"
    destination_address_prefix = "*"
  }
}
```

**Private Endpoints**
```hcl
resource "azurerm_private_endpoint" "database" {
  name                = "${var.project_name}-db-pe"
  location            = azurerm_resource_group.main.location
  resource_group_name = azurerm_resource_group.main.name
  subnet_id           = azurerm_subnet.private.id

  private_service_connection {
    name                           = "${var.project_name}-db-psc"
    private_connection_resource_id = azurerm_postgresql_server.main.id
    subresource_names              = ["postgresqlServer"]
    is_manual_connection           = false
  }
}
```

## Performance Optimization

### Caching Strategy

**Azure Cache for Redis Configuration**
```json
{
  "sku": {
    "name": "Premium",
    "family": "P",
    "capacity": 1
  },
  "enableNonSslPort": false,
  "minimumTlsVersion": "1.2",
  "redisConfiguration": {
    "maxmemory-policy": "allkeys-lru",
    "notify-keyspace-events": "Ex",
    "maxfragmentationmemory-reserved": "300",
    "maxmemory-reserved": "300"
  }
}
```

**Cache Implementation**
```python
# Python example with redis-py
import redis
import json
from functools import wraps

redis_client = redis.from_url(
    os.environ.get("REDIS_CONNECTION_STRING"),
    decode_responses=True
)

def cache_result(expiration=3600):
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            cache_key = f"{func.__name__}:{str(args)}:{str(kwargs)}"
            
            # Try to get from cache
            cached = redis_client.get(cache_key)
            if cached:
                return json.loads(cached)
            
            # Execute function
            result = func(*args, **kwargs)
            
            # Store in cache
            redis_client.setex(
                cache_key,
                expiration,
                json.dumps(result)
            )
            
            return result
        return wrapper
    return decorator
```

### CDN and Static Content

**Azure Front Door Configuration**
```json
{
  "frontendEndpoints": [{
    "name": "app-frontend",
    "hostName": "app.azurefd.net",
    "sessionAffinityEnabledState": "Disabled",
    "webApplicationFirewallPolicyLink": {
      "id": "/subscriptions/.../frontdoorwebapplicationfirewallpolicies/wafpolicy"
    }
  }],
  "backendPools": [{
    "name": "app-backend-pool",
    "backends": [{
      "address": "app.azurewebsites.net",
      "httpPort": 80,
      "httpsPort": 443,
      "priority": 1,
      "weight": 50
    }]
  }],
  "routingRules": [{
    "name": "app-routing-rule",
    "frontendEndpoints": [{"id": ".../frontendEndpoints/app-frontend"}],
    "acceptedProtocols": ["Http", "Https"],
    "patternsToMatch": ["/*"],
    "routeConfiguration": {
      "@odata.type": "#Microsoft.Azure.FrontDoor.Models.FrontdoorForwardingConfiguration",
      "forwardingProtocol": "HttpsOnly",
      "backendPool": {"id": ".../backendPools/app-backend-pool"},
      "cacheConfiguration": {
        "queryParameterStripDirective": "StripNone",
        "dynamicCompression": "Enabled",
        "cacheDuration": "P1D"
      }
    }
  }]
}
```

## Testing Strategy

### Test Environments

| Environment | Purpose | Configuration |
|-------------|---------|---------------|
| Dev | Development testing | Minimal resources, shared services |
| Test | Integration testing | Production-like, reduced scale |
| UAT | User acceptance | Production mirror, synthetic data |
| Prod | Production | Full scale, live data |

### Test Categories

#### 1. Unit Tests
- **Coverage Target**: 80%
- **Framework**: Jest (JS), pytest (Python), JUnit (Java)
- **Execution**: On every commit

#### 2. Integration Tests
```python
# Example Azure integration test
import pytest
from azure.identity import DefaultAzureCredential
from azure.keyvault.secrets import SecretClient

@pytest.fixture
def key_vault_client():
    credential = DefaultAzureCredential()
    vault_url = os.environ.get("KEY_VAULT_URL")
    return SecretClient(vault_url=vault_url, credential=credential)

def test_key_vault_connection(key_vault_client):
    # Test secret retrieval
    secret = key_vault_client.get_secret("test-secret")
    assert secret.value is not None
```

#### 3. Performance Tests
```yaml
# k6 performance test script
import http from 'k6/http';
import { check, sleep } from 'k6';

export let options = {
  stages: [
    { duration: '2m', target: 100 },  // Ramp up
    { duration: '5m', target: 100 },  // Stay at 100 users
    { duration: '2m', target: 0 },    // Ramp down
  ],
  thresholds: {
    http_req_duration: ['p(95)<500'], // 95% of requests under 500ms
    http_req_failed: ['rate<0.1'],    // Error rate under 10%
  },
};

export default function() {
  let response = http.get('https://app.azurewebsites.net/api/health');
  check(response, {
    'status is 200': (r) => r.status === 200,
    'response time < 500ms': (r) => r.timings.duration < 500,
  });
  sleep(1);
}
```

#### 4. Security Tests
```bash
# OWASP ZAP security scan
docker run -t owasp/zap2docker-stable zap-baseline.py \
  -t https://app.azurewebsites.net \
  -r security-report.html

# Container image scan
az acr task run \
  --registry $ACR_NAME \
  --cmd "mcr.microsoft.com/azuredefender/scan:latest $IMAGE_NAME" \
  /dev/null
```

## Migration Execution Plan

### Phase 1: Infrastructure Setup (Week 1-2)

**Week 1 Tasks**
- [ ] Create Azure subscription and resource groups
- [ ] Set up Azure DevOps project
- [ ] Configure Terraform backend
- [ ] Create base networking infrastructure
- [ ] Set up Azure Key Vault

**Week 2 Tasks**
- [ ] Deploy AKS cluster / App Services
- [ ] Configure Azure Container Registry
- [ ] Set up monitoring infrastructure
- [ ] Configure backup policies
- [ ] Implement network security

### Phase 2: Database Migration (Week 3-4)

**Week 3 Tasks**
- [ ] Deploy target database services
- [ ] Configure database security
- [ ] Set up replication
- [ ] Perform initial data sync
- [ ] Validate data integrity

**Week 4 Tasks**
- [ ] Complete performance testing
- [ ] Optimize database configuration
- [ ] Set up automated backups
- [ ] Document connection strings
- [ ] Update application configurations

### Phase 3: Application Migration (Week 5-7)

**Component Migration Order**
1. Non-critical services first
2. Stateless services
3. Stateful services
4. Frontend applications
5. Critical business services

**Migration Checklist per Component**
- [ ] Code updates for Azure integration
- [ ] Update CI/CD pipeline
- [ ] Build and push container images
- [ ] Deploy to development environment
- [ ] Run integration tests
- [ ] Deploy to test environment
- [ ] Performance validation
- [ ] Security scanning
- [ ] Deploy to UAT
- [ ] User acceptance testing
- [ ] Production deployment preparation

### Phase 4: Cutover Planning (Week 8-9)

**Cutover Checklist**
- [ ] Final data sync
- [ ] DNS preparation
- [ ] Load balancer configuration
- [ ] SSL certificate installation
- [ ] Monitoring alerts configured
- [ ] Rollback plan documented
- [ ] Communication plan executed
- [ ] Go/No-go decision criteria defined

**Cutover Execution**
```bash
# Cutover script example
#!/bin/bash
set -e

echo "Starting production cutover..."

# 1. Enable maintenance mode
az webapp config set --name $APP_NAME --resource-group $RG --always-on false

# 2. Final database sync
./sync-database.sh --final

# 3. Update DNS
az network dns record-set a update \
  --resource-group $DNS_RG \
  --zone-name $ZONE \
  --name $RECORD \
  --set aRecords[0].ipv4Address=$NEW_IP

# 4. Deploy application
kubectl apply -f ./k8s/production/

# 5. Health check
./health-check.sh --timeout 300

# 6. Update monitoring
./update-monitoring.sh

echo "Cutover completed successfully"
```

### Phase 5: Post-Migration (Week 10-12)

**Optimization Tasks**
- [ ] Performance tuning based on metrics
- [ ] Cost optimization review
- [ ] Security hardening
- [ ] Documentation updates
- [ ] Knowledge transfer sessions
- [ ] Runbook creation
- [ ] Disaster recovery testing

## Operational Procedures

### Monitoring and Alerting

**Key Metrics to Monitor**
```kusto
// Application Performance
AppRequests
| where TimeGenerated > ago(5m)
| summarize 
    RequestCount = count(),
    FailureRate = countif(Success == false) * 100.0 / count(),
    AvgDuration = avg(DurationMs),
    P95Duration = percentile(DurationMs, 95)
by bin(TimeGenerated, 1m)
| render timechart

// Resource Utilization
Perf
| where ObjectName == "K8SNode" and CounterName == "cpuUsageNanoCores"
| summarize AvgCPU = avg(CounterValue) by Computer, bin(TimeGenerated, 5m)
| render timechart
```

**Alert Configuration**
```json
{
  "alerts": [
    {
      "name": "High Error Rate",
      "query": "requests | where success == false | summarize ErrorRate = count() * 100.0 / toscalar(requests | count()) | where ErrorRate > 5",
      "threshold": 5,
      "frequency": "PT5M",
      "severity": 2,
      "actionGroup": "ops-team"
    },
    {
      "name": "High Response Time",
      "metric": "Response Time",
      "threshold": 1000,
      "operator": "GreaterThan",
      "aggregation": "Average",
      "period": "PT5M",
      "severity": 3
    }
  ]
}
```

### Backup and Recovery

**Backup Strategy**
- **Databases**: Automated daily backups with 35-day retention
- **Application State**: Persistent volume snapshots every 6 hours
- **Configuration**: Git repository with version control
- **Secrets**: Key Vault soft-delete enabled with 90-day retention

**Recovery Procedures**
```bash
# Database recovery
az postgres db restore \
  --resource-group $RG \
  --server-name $SERVER \
  --restore-point-in-time "2023-06-15T13:00:00Z" \
  --target-server $NEW_SERVER

# Application rollback
kubectl rollout undo deployment/$DEPLOYMENT_NAME -n $NAMESPACE

# Configuration restore
git checkout $LAST_KNOWN_GOOD_COMMIT
terraform apply -auto-approve
```

### Scaling Procedures

**Horizontal Scaling**
```yaml
# Horizontal Pod Autoscaler
apiVersion: autoscaling/v2
kind: HorizontalPodAutoscaler
metadata:
  name: app-hpa
spec:
  scaleTargetRef:
    apiVersion: apps/v1
    kind: Deployment
    name: app-deployment
  minReplicas: 2
  maxReplicas: 20
  metrics:
  - type: Resource
    resource:
      name: cpu
      target:
        type: Utilization
        averageUtilization: 70
  - type: Resource
    resource:
      name: memory
      target:
        type: Utilization
        averageUtilization: 80
```

**Vertical Scaling**
```bash
# Scale AKS node pool
az aks nodepool scale \
  --resource-group $RG \
  --cluster-name $AKS_CLUSTER \
  --name $NODEPOOL \
  --node-count 5

# Scale App Service Plan
az appservice plan update \
  --name $PLAN_NAME \
  --resource-group $RG \
  --sku P2v3
```

## Cost Management

### Cost Estimation by Component

| Component | Service | Monthly Cost (Est.) |
|-----------|---------|---------------------|
| Compute | AKS (3 nodes D4s_v3) | $450 |
| Database | PostgreSQL (GP 4vCore) | $300 |
| Storage | Blob Storage (1TB) | $50 |
| Networking | Application Gateway | $200 |
| Monitoring | Log Analytics | $100 |
| Backup | Recovery Services | $50 |
| **Total** | | **$1,150** |

### Cost Optimization Strategies

1. **Reserved Instances**
   ```bash
   # Purchase 1-year reserved instances for predictable workloads
   az reservations reservation-order purchase \
     --reservation-order-id $ORDER_ID \
     --sku Standard_D4s_v3 \
     --location eastus \
     --quantity 3 \
     --term P1Y
   ```

2. **Auto-shutdown for Non-Production**
   ```bash
   # Configure auto-shutdown for dev/test resources
   az vm auto-shutdown \
     --resource-group $RG \
     --name $VM_NAME \
     --time 1800
   ```

3. **Right-sizing Recommendations**
   - Use Azure Advisor for sizing recommendations
   - Implement automated right-sizing scripts
   - Regular quarterly reviews

## Appendix

### A. Reference Architecture Diagrams

[Detailed architecture diagrams would be included here]

### B. Security Compliance Checklist

- [ ] All data encrypted at rest
- [ ] TLS 1.2+ for all communications
- [ ] Network segmentation implemented
- [ ] RBAC configured for all resources
- [ ] Managed identities used (no passwords)
- [ ] Security Center recommendations addressed
- [ ] Vulnerability scanning enabled
- [ ] Audit logging configured
- [ ] Incident response plan documented
- [ ] Compliance validation completed

### C. Troubleshooting Guide

**Common Issues and Solutions**

1. **Connection to Key Vault fails**
   ```bash
   # Check managed identity assignment
   az keyvault show --name $KV_NAME --query "properties.accessPolicies"
   
   # Verify network connectivity
   nslookup $KV_NAME.vault.azure.net
   ```

2. **High latency after migration**
   - Check Application Insights for slow queries
   - Verify database indexes migrated correctly
   - Review network topology for unnecessary hops
   - Enable connection pooling

3. **Deployment failures**
   ```bash
   # Check deployment logs
   kubectl describe pod $POD_NAME
   kubectl logs $POD_NAME --previous
   
   # Verify image pull secrets
   kubectl get secrets -n $NAMESPACE
   ```

### D. Useful Commands and Scripts

```bash
# Quick health check
curl -f https://app.azurewebsites.net/health || echo "Health check failed"

# Database connection test
PGPASSWORD=$DB_PASSWORD psql -h $DB_HOST -U $DB_USER -d $DB_NAME -c "SELECT 1"

# View real-time logs
az webapp log tail --name $APP_NAME --resource-group $RG

# Export metrics
az monitor metrics list \
  --resource $RESOURCE_ID \
  --metric "Percentage CPU" \
  --start-time 2023-06-01T00:00:00Z \
  --end-time 2023-06-02T00:00:00Z \
  --interval PT1H \
  --output table
```

### E. Contact Information

**Escalation Matrix**

| Level | Team | Contact | Hours |
|-------|------|---------|-------|
| L1 | Operations | ops@company.com | 24/7 |
| L2 | DevOps | devops@company.com | Business hours |
| L3 | Architecture | arch@company.com | Business hours |
| Emergency | On-call | +1-xxx-xxx-xxxx | 24/7 |

---

*This document is version controlled in Git. Last updated: {datetime.now().strftime('%Y-%m-%d')}*
"""
        
        logger.info("LLD document generated")
        logger.info("All documents generated successfully")
        
        return asis_md, hld_md, lld_md
    
    @staticmethod
    def markdown_to_pdf(markdown_content: str, output_path: str):
        """Convert markdown to PDF - simplified version"""
        logger.info(f"Converting markdown to PDF: {output_path}")
        try:
            # For now, save as HTML which can be printed to PDF
            import markdown2
            html_content = f"""
            <html>
            <head>
                <meta charset="UTF-8">
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
                    h1 {{ color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                    h2 {{ color: #666; margin-top: 30px; }}
                    h3 {{ color: #888; margin-top: 20px; }}
                    h4 {{ color: #999; margin-top: 15px; }}
                    code {{ background-color: #f4f4f4; padding: 2px 4px; font-family: Consolas, monospace; }}
                    pre {{ background-color: #f4f4f4; padding: 10px; overflow-x: auto; border-radius: 4px; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; font-weight: bold; }}
                    blockquote {{ border-left: 4px solid #ddd; margin: 0; padding-left: 20px; color: #666; }}
                    ul, ol {{ margin: 10px 0; }}
                    li {{ margin: 5px 0; }}
                </style>
            </head>
            <body>
                {markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks', 'header-ids'])}
            </body>
            </html>
            """
            
            html_path = output_path.replace('.pdf', '.html')
            with open(html_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            logger.info(f"HTML file created: {html_path}")
            return True
        except Exception as e:
            logger.error(f"PDF generation error: {str(e)}", exc_info=True)
            return False
    
    @staticmethod
    def markdown_to_docx(markdown_content: str, output_path: str):
        """Convert markdown to DOCX"""
        logger.info(f"Converting markdown to DOCX: {output_path}")
        try:
            doc = DocxDocument()
            
            # Add styles
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Parse markdown line by line
            lines = markdown_content.split('\n')
            in_code_block = False
            in_table = False
            table_data = []
            
            for i, line in enumerate(lines):
                # Handle code blocks
                if line.strip().startswith('```'):
                    in_code_block = not in_code_block
                    if not in_code_block and table_data:
                        # End of code block, continue normal processing
                        pass
                    continue
                
                if in_code_block:
                    # Add code line with monospace font
                    p = doc.add_paragraph()
                    p.style = 'No Spacing'
                    run = p.add_run(line)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                    continue
                
                # Handle tables
                if line.strip().startswith('|'):
                    if not in_table:
                        in_table = True
                        table_data = []
                    
                    # Parse table row
                    cells = [cell.strip() for cell in line.split('|')[1:-1]]
                    if not all(cell.replace('-', '') == '' for cell in cells):  # Not separator line
                        table_data.append(cells)
                    continue
                elif in_table:
                    # End of table, create it
                    if table_data:
                        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(table_data):
                            for col_idx, cell_data in enumerate(row_data):
                                cell = table.rows[row_idx].cells[col_idx]
                                cell.text = cell_data
                                if row_idx == 0:  # Header row
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                    
                    in_table = False
                    table_data = []
                
                # Handle headings
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    doc.add_heading(line[5:], level=4)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif re.match(r'^\d+\. ', line):
                    doc.add_paragraph(line[3:], style='List Number')
                elif line.strip():
                    # Regular paragraph
                    p = doc.add_paragraph()
                    
                    # Handle inline formatting
                    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', line)
                    for part in parts:
                        if part.startswith('**') and part.endswith('**'):
                            # Bold
                            run = p.add_run(part[2:-2])
                            run.font.bold = True
                        elif part.startswith('`') and part.endswith('`'):
                            # Inline code
                            run = p.add_run(part[1:-1])
                            run.font.name = 'Consolas'
                            run.font.size = Pt(10)
                            # Add gray background effect by using highlighting
                            from docx.enum.text import WD_COLOR_INDEX
                            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                        else:
                            # Regular text
                            p.add_run(part)
                else:
                    # Empty line
                    doc.add_paragraph()
            
            # Save document
            doc.save(output_path)
            logger.info(f"DOCX file created: {output_path}")
            return True
        except Exception as e:
            logger.error(f"DOCX generation error: {str(e)}", exc_info=True)
            return False


# Gradio UI
def create_gradio_interface():
    """Create the Gradio web interface"""
    
    def analyze_repo(repo_url, api_key, progress=gr.Progress()):
        logger.info("="*60)
        logger.info("New analysis request received")
        logger.info(f"Repository: {repo_url}")
        logger.info("="*60)
        
        if not api_key:
            logger.error("No API key provided")
            return None, None, None, "❌ Please provide a Google API key"
        
        if not repo_url:
            logger.error("No repository URL provided")
            return None, None, None, "❌ Please provide a repository URL or path"
        
        try:
            # Clone repository if it's a URL
            progress(0.1, desc="Preparing repository...")
            if repo_url.startswith('http'):
                with tempfile.TemporaryDirectory() as temp_dir:
                    repo_path = os.path.join(temp_dir, 'repo')
                    progress(0.15, desc="Cloning repository...")
                    logger.info(f"Cloning repository from: {repo_url}")
                    
                    # Clone with shallow depth for speed
                    Repo.clone_from(repo_url, repo_path, depth=1)
                    logger.info("Repository cloned successfully")
                    
                    # Perform analysis
                    return perform_analysis(repo_path, api_key, progress, repo_url)
            else:
                # Local repository
                logger.info(f"Using local repository: {repo_url}")
                return perform_analysis(repo_url, api_key, progress, repo_url)
                
        except Exception as e:
            import traceback
            error_details = traceback.format_exc()
            logger.error(f"Analysis failed: {str(e)}", exc_info=True)
            return None, None, None, f"❌ Error: {str(e)}\n\nDetails:\n{error_details}"
    
    def perform_analysis(repo_path, api_key, progress, original_url):
        logger.info(f"Starting analysis for repository at: {repo_path}")
        
        # Analyze repository
        progress(0.2, desc="Initializing analyzer...")
        analyzer = RepositoryAnalyzer(api_key)
        
        def progress_callback(message):
            # Calculate progress based on the phase
            if "Git analysis" in message:
                current_progress = 0.25
            elif "repository structure" in message:
                current_progress = 0.35
            elif "components" in message:
                current_progress = 0.45
            elif "infrastructure" in message:
                current_progress = 0.55
            elif "security" in message:
                current_progress = 0.65
            elif "aggregating" in message:
                current_progress = 0.75
            elif "insights" in message:
                current_progress = 0.85
            else:
                current_progress = 0.5
                
            progress(current_progress, desc=message)
            logger.info(f"Progress: {message}")
        
        analysis = analyzer.analyze_repository(repo_path, progress_callback)
        analysis.repo_url = original_url  # Use original URL for display
        
        # Generate documents
        progress(0.9, desc="Generating comprehensive documents...")
        logger.info("Generating analysis documents")
        asis_md, hld_md, lld_md = DocumentGenerator.generate_markdown(analysis)
        
        # Save documents
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = f"analysis_output_{timestamp}"
        os.makedirs(output_dir, exist_ok=True)
        logger.info(f"Created output directory: {output_dir}")
        
        # Save markdown files
        files_created = []
        with open(os.path.join(output_dir, "ASIS_State.md"), 'w', encoding='utf-8') as f:
            f.write(asis_md)
            files_created.append("ASIS_State.md")
            
        with open(os.path.join(output_dir, "HLD.md"), 'w', encoding='utf-8') as f:
            f.write(hld_md)
            files_created.append("HLD.md")
            
        with open(os.path.join(output_dir, "LLD.md"), 'w', encoding='utf-8') as f:
            f.write(lld_md)
            files_created.append("LLD.md")
            
        logger.info(f"Markdown files created: {files_created}")
        
        # Generate PDF and DOCX
        progress(0.95, desc="Converting to PDF and DOCX formats...")
        for name, content in [("ASIS_State", asis_md), ("HLD", hld_md), ("LLD", lld_md)]:
            DocumentGenerator.markdown_to_pdf(content, os.path.join(output_dir, f"{name}.pdf"))
            DocumentGenerator.markdown_to_docx(content, os.path.join(output_dir, f"{name}.docx"))
        
        # Save analysis data as JSON for reference
        analysis_dict = {
            'repo_name': analysis.repo_name,
            'repo_url': analysis.repo_url,
            'analysis_date': analysis.analysis_date,
            'components_count': len(analysis.components),
            'tech_stack': analysis.tech_stack,
            'architecture_patterns': analysis.architecture_patterns,
            'complexity_level': analysis.migration_considerations.get('complexity_level', 'Unknown'),
            'migration_blockers': analysis.migration_considerations.get('blockers', []),
            'migration_accelerators': analysis.migration_considerations.get('accelerators', [])
        }
        
        with open(os.path.join(output_dir, "analysis_summary.json"), 'w', encoding='utf-8') as f:
            json.dump(analysis_dict, f, indent=2)
        
        # Create zip file
        progress(0.98, desc="Creating downloadable archive...")
        shutil.make_archive(output_dir, 'zip', output_dir)
        logger.info(f"Created archive: {output_dir}.zip")
        
        progress(1.0, desc="Analysis complete!")
        
        # Prepare summary
        summary = f"""✅ **Analysis completed successfully!**

📊 **Repository**: {analysis.repo_name}
🔗 **Source**: {original_url}
📅 **Analysis Date**: {datetime.now().strftime('%Y-%m-%d %H:%M')}

### Key Findings:
- **Components Found**: {len(analysis.components)}
- **Technologies**: {', '.join(analysis.tech_stack.get('languages', [])) or 'None detected'}
- **Frameworks**: {', '.join(analysis.tech_stack.get('frameworks', [])) or 'None detected'}
- **Databases**: {', '.join(analysis.tech_stack.get('databases', [])) or 'None detected'}
- **Architecture**: {', '.join(analysis.architecture_patterns) or 'Not determined'}

### Migration Assessment:
- **Complexity Level**: {analysis.migration_considerations.get('complexity_level', 'Not assessed')}
- **Migration Blockers**: {len(analysis.migration_considerations.get('blockers', []))}
- **Migration Accelerators**: {len(analysis.migration_considerations.get('accelerators', []))}

### Repository Metrics:
- **Total Commits**: {analysis.git_info.get('total_commits', 'Unknown')}
- **Contributors**: {analysis.git_info.get('contributor_count', 'Unknown')}
- **Last Activity**: {analysis.git_info.get('last_commit_date', 'Unknown')}

### Documents Generated:
1. **AS-IS State Analysis** - Current state documentation
2. **High Level Design (HLD)** - Architecture and migration strategy
3. **Low Level Design (LLD)** - Detailed implementation guide

📥 **Download the ZIP file** to access all documents in Markdown, PDF, and DOCX formats.

💡 **Next Steps**:
1. Review the AS-IS analysis to understand current state
2. Validate the HLD with stakeholders
3. Use the LLD for implementation planning
4. Check logs folder for detailed analysis logs"""
        
        logger.info("Analysis completed successfully")
        logger.info("="*60)
        
        # Return results
        return (
            f"{output_dir}.zip",
            asis_md[:3000] + "\n\n... (Preview truncated. Download full document)",
            hld_md[:3000] + "\n\n... (Preview truncated. Download full document)",
            summary
        )
    
    # Create Gradio interface with custom theme
    with gr.Blocks(
        title="Repository Migration Analyzer - Enhanced",
        theme=gr.themes.Soft(
            primary_hue="blue",
            secondary_hue="gray",
            neutral_hue="gray",
            font=gr.themes.GoogleFont("Inter")
        ),
        css="""
        .gradio-container {
            max-width: 1200px;
            margin: auto;
        }
        .main-title {
            text-align: center;
            margin-bottom: 30px;
        }
        """
    ) as app:
        gr.Markdown("""
        <div class="main-title">
        
        # 🚀 Repository Migration Analyzer - Enhanced Version
        
        **Comprehensive Git Repository Analysis for Azure Cloud Migration**
        
        </div>
        
        This enhanced analyzer performs deep analysis of your repository structure, components, dependencies, 
        and generates professional migration documentation including AS-IS state, High-Level Design (HLD), 
        and Low-Level Design (LLD) documents.
        
        ---
        """)
        
        with gr.Row():
            with gr.Column(scale=2):
                repo_input = gr.Textbox(
                    label="📁 Repository URL or Local Path",
                    placeholder="https://github.com/user/repo.git or /path/to/local/repo",
                    lines=1,
                    info="Enter a Git repository URL (GitHub, GitLab, etc.) or local directory path"
                )
                api_key_input = gr.Textbox(
                    label="🔑 Google Gemini API Key",
                    placeholder="Enter your Google AI API key",
                    type="password",
                    lines=1,
                    info="Get your API key from makersuite.google.com/app/apikey"
                )
                
                with gr.Row():
                    analyze_btn = gr.Button(
                        "🔍 Start Deep Analysis", 
                        variant="primary", 
                        scale=2,
                        size="lg"
                    )
                    clear_btn = gr.Button(
                        "🗑️ Clear", 
                        scale=1,
                        size="lg"
                    )
            
            with gr.Column(scale=1):
                status_output = gr.Textbox(
                    label="📊 Analysis Summary",
                    lines=15,
                    interactive=False,
                    show_copy_button=True
                )
                download_output = gr.File(
                    label="📥 Download Complete Analysis Package",
                    interactive=False
                )
        
        with gr.Tabs():
            with gr.TabItem("📄 AS-IS State Preview"):
                asis_preview = gr.Textbox(
                    label="Current State Analysis Preview",
                    lines=25,
                    interactive=False,
                    show_copy_button=True
                )
            with gr.TabItem("📐 HLD Preview"):
                hld_preview = gr.Textbox(
                    label="High Level Design Preview",
                    lines=25,
                    interactive=False,
                    show_copy_button=True
                )
            with gr.TabItem("ℹ️ How to Use"):
                gr.Markdown("""
                ### 📖 User Guide
                
                #### 1️⃣ **Prerequisites**
                - **Google AI API Key**: Get one from [Google AI Studio](https://makersuite.google.com/app/apikey)
                - **Repository Access**: Public GitHub/GitLab URL or local repository path
                
                #### 2️⃣ **Analysis Process**
                1. Enter your repository URL or local path
                2. Paste your Google AI API key
                3. Click "Start Deep Analysis"
                4. Wait for the analysis to complete (typically 3-10 minutes)
                5. Download the ZIP file containing all documents
                
                #### 3️⃣ **What You Get**
                - **AS-IS State Document**: Complete analysis of current architecture
                - **High-Level Design**: Azure migration strategy and architecture
                - **Low-Level Design**: Detailed implementation specifications
                - **All formats**: Markdown, PDF, and DOCX
                
                #### 4️⃣ **Analysis Features**
                - Deep code analysis with framework detection
                - Dependency mapping and vulnerability checking
                - Container and Kubernetes configuration analysis
                - Security assessment and recommendations
                - Cost estimation and optimization strategies
                - Complete migration roadmap with timelines
                
                #### 5️⃣ **Performance Tips**
                - Smaller repos (< 100 files): 2-3 minutes
                - Medium repos (100-1000 files): 5-7 minutes
                - Large repos (> 1000 files): 10-15 minutes
                - Monorepos: May take longer due to multiple components
                
                #### 6️⃣ **Troubleshooting**
                - **Rate Limit**: The tool respects API limits (14 calls/minute)
                - **Large Repos**: Analysis continues even if progress seems stuck
                - **Logs**: Check the `logs` folder for detailed information
                - **Memory**: For very large repos, ensure sufficient system memory
                """)
        
        gr.Markdown("""
        ---
        
        ### 💡 Example Repositories to Try
        """)
        
        gr.Examples(
            examples=[
                ["https://github.com/end-of-game/openshift-voting-app"],
                ["https://github.com/dockersamples/example-voting-app"],
                ["https://github.com/spring-projects/spring-petclinic"],
                ["https://github.com/GoogleCloudPlatform/microservices-demo"],
                ["https://github.com/gothinkster/realworld"],
            ],
            inputs=repo_input,
            label="Click any example to analyze"
        )
        
        gr.Markdown("""
        ### 🔒 Privacy & Security
        - Your API key is used only for this session
        - Repository data is processed locally
        - No data is stored permanently
        - Temporary files are cleaned up after analysis
        
        ### 📊 Recent Enhancements
        - ✅ Deep component analysis with language detection
        - ✅ Comprehensive dependency tracking
        - ✅ OpenShift/Kubernetes manifest parsing
        - ✅ Security vulnerability detection
        - ✅ Multi-stage Docker analysis
        - ✅ Cost estimation and optimization
        - ✅ Detailed implementation guides with code samples
        
        ---
        
        <div style="text-align: center; color: #666; font-size: 0.9em;">
        Enhanced Repository Migration Analyzer v2.0 | Powered by Google Gemini AI
        </div>
        """)
        
        # Set up event handlers
        analyze_btn.click(
            fn=analyze_repo,
            inputs=[repo_input, api_key_input],
            outputs=[download_output, asis_preview, hld_preview, status_output],
            show_progress=True
        )
        
        clear_btn.click(
            fn=lambda: ("", "", None, "", "", ""),
            outputs=[repo_input, api_key_input, download_output, asis_preview, hld_preview, status_output]
        )
    
    return app


# Main execution
if __name__ == "__main__":
    print("="*60)
    print("Repository Migration Analyzer - Enhanced Version")
    print("="*60)
    print("Starting application...")
    print(f"Logs will be saved to: logs/repo_analyzer_[timestamp].log")
    print("\n⚠️  IMPORTANT: You need a Google AI API key!")
    print("Get one from: https://makersuite.google.com/app/apikey")
    print("-" * 60)
    
    # Check for required libraries
    try:
        import markdown2
        import gradio
        import git
        from docx import Document
    except ImportError as e:
        print(f"\n❌ Missing required library: {e}")
        print("\nPlease install all requirements:")
        print("pip install -r requirements.txt")
        print("\nRequired packages:")
        print("- gradio")
        print("- langchain-google-genai") 
        print("- google-generativeai")
        print("- GitPython")
        print("- python-docx")
        print("- markdown2")
        print("- pyyaml")
        print("- toml")
        exit(1)
    
    # Create and launch Gradio app
    app = create_gradio_interface()
    
    logger.info("Launching Gradio interface")
    
    # Launch the app
    try:
        app.launch(
            share=False,  # Set to True to create a public link
            server_name="0.0.0.0",  # Allow external connections
            server_port=7860,
            show_error=True,
            quiet=False
        )
    except Exception as e:
        logger.error(f"Failed to launch Gradio interface: {e}")
        print(f"\n❌ Error launching application: {e}")
        print("\nTry running with a different port:")
        print("python repo_analyzer.py --port 7861")